"""Document primitives for the CMP7200 dissertation.

Formatting follows the assignment brief: font size 11, 1.5 line spacing, so
markers have room to annotate. Everything before the introduction and
everything after the conclusion is excluded from the word count, so the
front matter and appendices are set apart deliberately.
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
LINE_SPACING = 1.5

INK = RGBColor(0x1F, 0x29, 0x33)
HEAD1 = RGBColor(0x1A, 0x36, 0x5D)
HEAD2 = RGBColor(0x2C, 0x52, 0x82)
HEAD3 = RGBColor(0x3D, 0x4F, 0x63)
GREY = RGBColor(0x5B, 0x6B, 0x7A)
RULE = "C3CCD6"


# ── Document setup ───────────────────────────────────────────────────────

AUTHOR = "Abdul Wahab"


def new_document() -> Document:
    doc = Document()

    # Word shows this under File > Properties. Left alone, python-docx stamps
    # its own name there, so every document produced by this project would
    # credit the library rather than the author.
    props = doc.core_properties
    props.author = AUTHOR
    props.last_modified_by = AUTHOR

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(8)
    # East-Asian font mapping, otherwise Word substitutes for some glyphs
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return doc


def add_page_numbers(doc: Document):
    """Page number in the footer, centred."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr, kind in (("begin", "fldChar"), ("PAGE", "instrText"), ("end", "fldChar")):
        el = OxmlElement(f"w:{kind}")
        if kind == "fldChar":
            el.set(qn("w:fldCharType"), instr)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run._r.append(el)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


# ── Headings ─────────────────────────────────────────────────────────────

def h1(doc, text, *, page_break=True):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(17)
        r.font.color.rgb = HEAD1
        r.font.bold = True
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(13)
        r.font.color.rgb = HEAD2
        r.font.bold = True
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(11.5)
        r.font.color.rgb = HEAD3
        r.font.bold = True
        r.font.italic = False
    return p


# ── Body text ────────────────────────────────────────────────────────────

def _style_run(run, *, bold=False, italic=False, size=None, colour=None):
    run.font.name = BODY_FONT
    run.font.size = size or BODY_SIZE
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour or INK
    return run


def para(doc, text="", *, bold=False, italic=False, size=None, colour=None,
         align="justify", space_after=8, indent=None):
    p = doc.add_paragraph()
    p.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        _style_run(p.add_run(text), bold=bold, italic=italic, size=size, colour=colour)
    return p


def bullet(doc, text, *, lead=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    if lead:
        _style_run(p.add_run(lead), bold=True)
    _style_run(p.add_run(text))
    return p


def numbered(doc, text, *, lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.75)
    if lead:
        _style_run(p.add_run(lead), bold=True)
    _style_run(p.add_run(text))
    return p


def quote(doc, text):
    """An indented, set-apart statement — used for the aim and the research question."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.35
    _style_run(p.add_run(text), italic=True, colour=HEAD2)
    _shade(p, "F2F5F9")
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = INK
    _shade(p, "F5F6F8")
    return p


def _shade(paragraph, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    paragraph._p.get_or_add_pPr().append(shd)


# ── Figures ──────────────────────────────────────────────────────────────

def figure(doc, path, caption, *, width_cm=15.5):
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(8)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
    p = para(doc, caption, size=Pt(9), colour=GREY, align="center",
             space_after=14, italic=True)
    return p


# ── Tables ───────────────────────────────────────────────────────────────

def table(doc, headers, rows, *, widths=None, caption=None, font_size=9.5,
          first_col_bold=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(str(text))
        r.font.bold = True
        r.font.size = Pt(font_size)
        r.font.name = BODY_FONT
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], "2C5282")

    for row_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
            r.font.name = BODY_FONT
            r.font.color.rgb = INK
            r.font.bold = first_col_bold and i == 0
        if row_i % 2 == 1:
            for c in cells:
                _shade_cell(c, "F4F6F9")

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)

    if caption:
        para(doc, caption, size=Pt(9), colour=GREY, align="center",
             space_after=14, italic=True)
    else:
        para(doc, space_after=6)
    return t


def _shade_cell(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


# ── Front matter ─────────────────────────────────────────────────────────

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc_entry(doc, text, page=None, *, level=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(level * 0.7)
    _style_run(p.add_run(text), bold=bold,
               size=Pt(11 if level == 0 else 10.5),
               colour=INK if level == 0 else GREY)
    return p
