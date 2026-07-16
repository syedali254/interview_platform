"""Generate PROJECT_GUIDE.docx with detailed file-by-file explanations."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper functions ──
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2D, 0x2D, 0x5E)

def h3(text):
    doc.add_heading(text, level=3)

def para(text):
    doc.add_paragraph(text)

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def bold_para(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    p.add_run(text)

def bullet(text, level=0):
    doc.add_paragraph(text, style='List Bullet')


# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
title = doc.add_heading('InterviewAI — Complete Project Guide', level=0)
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

para("A complete explanation of every file, how it works, and how everything fits together.")
para("Written for developers who want to understand the codebase quickly.")
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    "1. What Is InterviewAI?",
    "2. How to Run the Project",
    "3. Project Folder Structure",
    "4. The Main File — app.py",
    "5. Core Settings — config.py",
    "6. AI Brain — llm.py & gemini.py",
    "7. CV Reader — cv_agent.py",
    "8. JD Reader — jd_agent.py",
    "9. Question Maker — question_agent.py",
    "10. Skill Graph — skill_graph.py",
    "11. Interview State Tracker — state.py",
    "12. Next Question Chooser — traversal.py",
    "13. Graph Drawer — visualize.py",
    "14. Answer Scorer — evaluator.py",
    "15. Interview Manager — interview_loop.py",
    "16. Report Maker — generator.py",
    "17. Interview Room UI — room_components.py",
    "18. Voice System (LiveKit) — Overview",
    "19. Voice Helpers — voice.py",
    "20. Voice Web Server — whisper_server.py",
    "21. Voice Launcher — launcher.py",
    "22. Voice Agent — run_agent.py",
    "23. Browser Client — client.html",
    "24. LiveKit Server Config — livekit.yaml",
    "25. Extra Files (Dead Code)",
    "26. API Keys & Security",
    "27. Data Flow Summary"
]
for item in toc_items:
    bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. WHAT IS INTERVIEWAI?
# ════════════════════════════════════════════════════════════
h1("1. What Is InterviewAI?")

para(
    "InterviewAI is a smart interview platform that helps companies conduct technical "
    "interviews automatically. You give it a candidate's CV (resume) and a job description, "
    "and it does the rest:"
)

bullet("Reads the CV and extracts skills, experience, education, projects")
bullet("Reads the job description and extracts required skills")
bullet("Builds a skill map from a European database of 1200+ skills (ESCO)")
bullet("Compares what the candidate has vs what the job needs → finds gaps")
bullet("Generates interview questions targeting those gaps")
bullet("Conducts a live interview (text, voice, or video call)")
bullet("Scores each answer using AI")
bullet("Generates a final report with hiring recommendation")

para(
    "The system is built in Python using Streamlit for the web interface, Google Gemini "
    "for AI reasoning, and optionally LiveKit for real-time voice/video interviews."
)

# ════════════════════════════════════════════════════════════
# 2. HOW TO RUN
# ════════════════════════════════════════════════════════════
h1("2. How to Run the Project")

para("Step 1 — Go to the project folder:")
code("cd InterviewAI")

para("Step 2 — Install required libraries (one-time):")
code("pip install -r requirements.txt")

para("Step 3 — Create a .env file with your API keys:")
code("GEMINI_API_KEY=your_gemini_key_here")
code("DEEPGRAM_API_KEY=your_deepgram_key_here")
code("ELEVENLABS_API_KEY=your_elevenlabs_key_here")

para("Step 4 — Run the app:")
code("streamlit run app.py")

para("Step 5 — Open the URL shown in terminal (usually http://localhost:8501)")

# ════════════════════════════════════════════════════════════
# 3. FOLDER STRUCTURE
# ════════════════════════════════════════════════════════════
h1("3. Project Folder Structure")

code("InterviewAI/")
code("├── app.py                       # Main file — run this")
code("├── requirements.txt             # List of Python libraries")
code("├── .env                         # API keys (DO NOT SHARE)")
code("├── .gitignore                   # Files to exclude from git")
code("├── README.md                    # Old documentation")
code("├── PROJECT_GUIDE.md             # This guide (Markdown)")
code("├──")
code("├── core/                        # All backend logic")
code("│   ├── config.py                # Settings & configuration")
code("│   ├── llm.py                   # AI brain (talks to Gemini/Ollama)")
code("│   ├── gemini.py                # Old Gemini-only AI client")
code("│   ├──")
code("│   ├── agents/                  # Smart AI modules")
code("│   │   ├── cv_agent.py          # Reads CV → extracts info")
code("│   │   ├── jd_agent.py          # Reads job ad → extracts info")
code("│   │   └── question_agent.py    # Creates interview questions")
code("│   ├──")
code("│   ├── graph/                   # Skill knowledge graph")
code("│   │   ├── skill_graph.py       # Builds skill map from ESCO data")
code("│   │   ├── state.py             # Tracks interview progress")
code("│   │   ├── traversal.py         # Picks next skill to ask")
code("│   │   └── visualize.py         # Draws skill graphs as images")
code("│   ├──")
code("│   ├── evaluator/")
code("│   │   └── evaluator.py         # Scores candidate answers")
code("│   ├── pipeline/")
code("│   │   └── interview_loop.py    # Runs the live interview")
code("│   ├── report/")
code("│   │   └── generator.py         # Creates final report")
code("│   ├── interview/")
code("│   │   └── room_components.py   # HTML templates for interview UI")
code("│   └──")
code("│   └── livekit/                 # Voice/video call system")
code("│       ├── voice.py             # Deepgram STT + ElevenLabs TTS wrappers")
code("│       ├── whisper_server.py    # Mini HTTP server (port 18765)")
code("│       ├── launcher.py          # Starts/stops LiveKit from app.py")
code("│       ├── run_agent.py         # Voice AI agent")
code("│       ├── client.html          # Browser page for video interview")
code("│       ├── start_livekit.py     # Standalone launch script")
code("│       ├── agent_standalone.py  # Old standalone agent (not used)")
code("│       ├── adapter.py           # Old adapter (not used)")
code("│       └── livekit.yaml         # LiveKit server configuration")
code("├──")
code("├── frontend/                    # Old page-based UI (not used anymore)")
code("├── data/                        # Sample data & ESCO CSV files")
code("└── tests/                       # Test scripts")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. APP.PY
# ════════════════════════════════════════════════════════════
h1("4. The Main File — app.py")

bold_para("File", "app.py")
bold_para("Purpose", "This is the entry point. You run this file to start the whole application.")
bold_para("How it works", "It uses Streamlit (a Python library that turns Python scripts into web pages). "
           "When you run 'streamlit run app.py', it opens a webpage in your browser with 7 tabs.")

para("The 7 tabs are:")

bullet("Tab 1 — Upload: You upload a CV (as PDF or paste text) and a job description. "
       "Click 'Run Full Pipeline' to process everything.")
bullet("Tab 2 — Analysis: Shows what the AI extracted from your CV and job description "
       "(skills, experience, education, etc.)")
bullet("Tab 3 — Skill Graph: Shows bar charts comparing what skills the candidate has "
       "vs what the job needs. Shows match percentage and skill gaps.")
bullet("Tab 4 — Questions: Generates a full set of interview questions based on the skill gaps. "
       "Shows opening, technical, behavioral, and closing questions.")
bullet("Tab 5 — Live Interview: The actual interview. Three modes:")
bullet("   • Text Mode — you type answers manually")
bullet("   • Voice Mode — you speak, browser converts to text, AI responds")
bullet("   • LiveKit Mode — full video-call-style interview with voice AI", level=1)
bullet("Tab 6 — Report: Shows the final scores, per-skill breakdown, and hiring decision.")
bullet("Tab 7 — Logs: Shows debug messages — useful for troubleshooting.")

para(
    "The app also starts a small background web server (whisper_server.py) on port 18765 "
    "when it launches, which handles voice features."
)

# ════════════════════════════════════════════════════════════
# 5. CONFIG.PY
# ════════════════════════════════════════════════════════════
h1("5. Core Settings — config.py")

bold_para("File", "core/config.py")
bold_para("Purpose", "Stores all configuration settings and loads API keys from the .env file.")
bold_para("How it works", "When any module needs a setting (like an API key or a score threshold), "
           "it imports config.py and reads the value. config.py uses python-dotenv to load "
           "the .env file into environment variables.")

para("Key settings in this file:")

bullet("GEMINI_API_KEY — key for Google's Gemini AI")
bullet("GEMINI_MODEL — which Gemini model to use (default: gemini-2.5-flash)")
bullet("LLM_PROVIDER — whether to use 'gemini' (cloud) or 'ollama' (local)")
bullet("MAX_QUESTIONS_PER_SESSION — max 3 questions per interview")
bullet("SCORE_STRONG_THRESHOLD = 70 — score above 70 is 'strong'")
bullet("SCORE_WEAK_THRESHOLD = 40 — score below 40 is 'weak/gap'")
bullet("LIVEKIT_URL, LIVEKIT_API_KEY, LIVEKIT_API_SECRET — LiveKit server connection details")

# ════════════════════════════════════════════════════════════
# 6. LLM.PY & GEMINI.PY
# ════════════════════════════════════════════════════════════
h1("6. AI Brain — llm.py & gemini.py")

h2("llm.py (main AI client)")
bold_para("File", "core/llm.py")
bold_para("Purpose", "This is the main AI communication module. Every AI call in the project goes through this file.")
bold_para("How it works", "It can talk to two AI backends:")

bullet("Gemini (Google cloud) — sends HTTP requests to Google's Gemini API")
bullet("Ollama (local) — sends HTTP requests to Ollama running on your own computer")

para("Key functions:")
bullet("call_llm(prompt) — sends any text prompt to the AI, gets text response back")
bullet("call_llm_json(prompt) — sends prompt asking for JSON, gets structured data back")

para(
    "It automatically picks which backend to use based on the LLM_PROVIDER setting in config.py. "
    "If you have a Gemini API key, it uses Gemini by default (faster, smarter). "
    "If not, it falls back to Ollama (slower, but free and runs offline)."
)

h2("gemini.py (legacy)")
bold_para("File", "core/gemini.py")
bold_para("Purpose", "Old version of the Gemini client. Kept for reference but NOT used by the main pipeline anymore.")
bold_para("Note", "Everything in this file has been replaced by llm.py. It's dead code.")

# ════════════════════════════════════════════════════════════
# 7. CV_AGENT.PY
# ════════════════════════════════════════════════════════════
h1("7. CV Reader — cv_agent.py")

bold_para("File", "core/agents/cv_agent.py")
bold_para("Purpose", "Reads a candidate's CV (resume) and extracts structured information from it.")
bold_para("How it works", "Two functions:")

bullet("parse_cv_text(text) — takes raw text, sends it to Gemini AI with a prompt that says "
       "'extract name, skills, experience, education, projects from this CV as JSON'. "
       "Returns structured JSON data.")
bullet("parse_cv_pdf(pdf_bytes) — takes a PDF file, extracts the text using PyMuPDF library, "
       "then passes the text to parse_cv_text().")

para(
    "The AI prompt is designed to understand different CV formats. It can handle bullet points, "
    "paragraphs, tables — whatever format the CV uses. The output is always structured JSON "
    "so other modules can easily work with it."
)

# ════════════════════════════════════════════════════════════
# 8. JD_AGENT.PY
# ════════════════════════════════════════════════════════════
h1("8. JD Reader — jd_agent.py")

bold_para("File", "core/agents/jd_agent.py")
bold_para("Purpose", "Reads a job description and extracts structured requirements from it.")
bold_para("How it works", "Similar to cv_agent.py. It takes the job description text and asks Gemini "
           "to extract: job title, required skills, nice-to-have skills, experience level, domain, "
           "and key responsibilities — all as structured JSON.")

# ════════════════════════════════════════════════════════════
# 9. QUESTION_AGENT.PY
# ════════════════════════════════════════════════════════════
h1("9. Question Maker — question_agent.py")

bold_para("File", "core/agents/question_agent.py")
bold_para("Purpose", "Generates interview questions based on the skill gaps found by the graph.")
bold_para("How it works", "Two modes:")

para("Mode 1 — Full Question Set (Tab 4):")
bullet("generate_interview_questions(topics, cv_data, jd_data) — takes the interview topics "
       "from the skill graph and generates a complete set: opening questions, technical questions "
       "(one per skill gap), behavioral questions, and closing questions.")

para("Mode 2 — Single Question (during live interview):")
bullet("generate_position_question(skill, difficulty, cv_data, jd_data) — generates exactly one "
       "question targeting a specific skill at a specific difficulty level. Used during the "
       "live interview when the system needs to ask about a particular skill.")

# ════════════════════════════════════════════════════════════
# 10. SKILL_GRAPH.PY
# ════════════════════════════════════════════════════════════
h1("10. Skill Knowledge Graph — skill_graph.py")

bold_para("File", "core/graph/skill_graph.py")
bold_para("Purpose", "Builds a map (graph) of 1200+ skills and compares candidate vs job skills.")
bold_para("How it works", "This is the smartest module. It:")

bullet("1. Loads ESCO data — a European database of 1200+ digital skills with their relationships "
       "(e.g., 'Python' is related to 'Programming', which is related to 'Software Development')")
bullet("2. Builds a NetworkX graph — each skill is a node, relationships are edges")
bullet("3. Matches CV skills to the graph using fuzzy matching (handles typos and variations)")
bullet("4. Matches job skills to the graph the same way")
bullet("5. Compares both → finds which job skills the candidate is missing")
bullet("6. Calculates match percentage")
bullet("7. Generates interview topics — prioritizes missing important skills first")

para("Key functions:")
bullet("match_skill(text) — finds the closest ESCO skill to any text")
bullet("add_candidate_skills(skills) / add_job_skills(required, nice_to_have)")
bullet("analyse_gaps() — returns match percentage, matched/missing skills")
bullet("get_interview_topics() — returns prioritized topics for the interview")

# ════════════════════════════════════════════════════════════
# 11. STATE.PY
# ════════════════════════════════════════════════════════════
h1("11. Interview State Tracker — state.py")

bold_para("File", "core/graph/state.py")
bold_para("Purpose", "Tracks the progress of each skill during the live interview.")
bold_para("How it works", "For every skill, it stores:")

bullet("Status — pending / verified_strong / verified_weak / confirmed_gap / skipped")
bullet("Score — the numeric score from the evaluator (0-100)")
bullet("Questions asked — how many questions were asked about this skill")
bullet("Best score — the highest score achieved")

para(
    "The InterviewState class provides helper methods like pending_skills (which skills still need "
    "to be asked), verified_strong_skills (which skills the candidate knows well), and "
    "confirmed_gaps (which skills the candidate is missing)."
)

# ════════════════════════════════════════════════════════════
# 12. TRAVERSAL.PY
# ════════════════════════════════════════════════════════════
h1("12. Next Question Chooser — traversal.py")

bold_para("File", "core/graph/traversal.py")
bold_para("Purpose", "Decides which skill to ask about next during the live interview.")
bold_para("How it works", "Four strategies:")

bullet("Adaptive (default) — picks the skill with the lowest score first. "
       "Weighted by priority. Avoids asking the same skill twice in a row.")
bullet("BFS — goes through skills one by one, round-robin style")
bullet("DFS — keeps asking about the same skill until it's verified, then moves on")
bullet("Spaced — picks the skill that has been asked the fewest questions")

para(
    "Also has decide_follow_up() — returns True if the candidate scored below the weak threshold "
    "and there are still follow-up questions remaining for this skill."
)

# ════════════════════════════════════════════════════════════
# 13. VISUALIZE.PY
# ════════════════════════════════════════════════════════════
h1("13. Graph Drawer — visualize.py")

bold_para("File", "core/graph/visualize.py")
bold_para("Purpose", "Creates PNG images of the skill graphs for display in the Streamlit UI.")
bold_para("How it works", "Uses matplotlib to draw 4 types of graphs:")

bullet("Candidate graph — shows skills the candidate has")
bullet("Job graph — shows skills the job requires")
bullet("Gap graph — highlights skills the candidate is missing")
bullet("Full graph — shows everything together")

# ════════════════════════════════════════════════════════════
# 14. EVALUATOR.PY
# ════════════════════════════════════════════════════════════
h1("14. Answer Scorer — evaluator.py")

bold_para("File", "core/evaluator/evaluator.py")
bold_para("Purpose", "Scores a candidate's answer to any interview question.")
bold_para("How it works", "Uses a technique called 'LLM-as-Judge' (the AI judges itself). Three steps:")

para("Step 1 — Generate Reference Answer:")
bullet("generate_reference_answer(question, skill) — asks Gemini to write what a good answer "
       "looks like (max 100 words). This is the 'ideal answer' to compare against.")

para("Step 2 — Score (Double Call):")
bullet("track_a_evaluate(question, answer, skill, reference) — calls Gemini TWICE to score "
       "the answer on 4 criteria, each worth 0-25 points:")
bullet("   • technical_accuracy — is the answer technically correct?", level=1)
bullet("   • completeness — did they cover all important points?", level=1)
bullet("   • clarity — is the answer well explained?", level=1)
bullet("   • relevance — is the answer on-topic?", level=1)
bullet("The two calls use shuffled criteria order to avoid bias. The final score is averaged.")

para("Step 3 — Verdict:")
bullet("evaluate_answer(question, answer, skill) — returns final_score (0-100), "
       "verdict ('strong' / 'weak' / 'gap'), and the reference answer.")

# ════════════════════════════════════════════════════════════
# 15. INTERVIEW_LOOP.PY
# ════════════════════════════════════════════════════════════
h1("15. Interview Manager — interview_loop.py")

bold_para("File", "core/pipeline/interview_loop.py")
bold_para("Purpose", "Conducts the entire live interview from start to finish.")
bold_para("How it works", "The InterviewLoop class is the main conductor. It connects all the pieces:")

bullet("Initialization — takes interview topics, CV data, job data. Creates the InterviewState.")
bullet("get_next_question() — uses traversal.py to pick the next skill, then question_agent.py "
       "to generate a question for that skill.")
bullet("submit_answer(question, answer_text) — takes the candidate's answer, sends it to "
       "evaluator.py for scoring, records the score in state.py, and decides whether a "
       "follow-up question is needed.")
bullet("is_complete — returns True when all skills are verified or max questions reached.")
bullet("get_summary() — returns the full state + question/answer history for the final report.")

# ════════════════════════════════════════════════════════════
# 16. GENERATOR.PY
# ════════════════════════════════════════════════════════════
h1("16. Report Maker — generator.py")

bold_para("File", "core/report/generator.py")
bold_para("Purpose", "Generates the final interview report with scores and recommendations.")
bold_para("How it works", "generate_report(loop) takes the completed InterviewLoop and produces:")

bullet("Overall score — weighted average of all skill scores")
bullet("Per-skill breakdown — each skill with its score and status")
bullet("Strengths — skills scored 70+")
bullet("Gaps — skills scored below 40 or confirmed missing")
bullet("Development areas — skills that need improvement (40-70)")
bullet("Answer log — every question asked and the candidate's answer")
bullet("Narrative summary — a paragraph explaining the result")
bullet("Hiring verdict — 'Strong Hire' / 'Weak Hire' / 'No Hire'")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 17. ROOM_COMPONENTS.PY
# ════════════════════════════════════════════════════════════
h1("17. Interview Room UI — room_components.py")

bold_para("File", "core/interview/room_components.py")
bold_para("Purpose", "Provides HTML and JavaScript templates for the interview page inside the app.")
bold_para("How it works", "Contains 3 main components:")

bullet("DEVICE_SETUP_HTML — a complete HTML page with mic test (shows audio level meter) "
       "and camera test (live preview). Used before the interview starts to make sure everything works.")
bullet("get_interview_question_html(question, q_number, ...) — a dark-themed card that displays "
       "the current question with a colored badge showing the type (technical/behavioral/etc.)")
bullet("get_controls_html(mic_on, cam_on) — a status bar showing mic/camera icons and recording indicator")

# ════════════════════════════════════════════════════════════
# 18. LIVEKIT OVERVIEW
# ════════════════════════════════════════════════════════════
h1("18. Voice System (LiveKit) — Overview")

para(
    "LiveKit is an open-source platform for real-time voice and video calls. In this project, "
    "we use it to conduct voice interviews where the AI speaks questions aloud and the candidate "
    "speaks answers naturally — like a real conversation."
)

h2("How the LiveKit Voice System Works (End to End)")

para("The LiveKit system has 3 main processes that all run at the same time:")

para("Process 1 — LiveKit Server:")
bullet("A binary program (livekit-server.exe) downloaded from LiveKit's GitHub releases")
bullet("Runs on your computer, listens on port 7880")
bullet("Acts like a switchboard — routes audio/video between participants (candidate + AI agent)")
bullet("Configuration is in livekit.yaml (dev key/secret for testing)")

para("Process 2 — Web Server (whisper_server.py):")
bullet("A small Python HTTP server running on port 18765")
bullet("Serves 3 purposes:")
bullet("   1. Creates room tokens — when a candidate joins, it generates a secure token", level=1)
bullet("   2. Serves the client webpage (client.html) at http://localhost:18765/livekit", level=1)
bullet("   3. Handles speech-to-text and text-to-speech requests", level=1)

para("Process 3 — Voice Agent (run_agent.py):")
bullet("A Python program that acts as the AI interviewer")
bullet("Connects to the LiveKit server and waits for a candidate to join")
bullet("When someone joins, it:")
bullet("   1. Speaks a greeting and first question (via ElevenLabs TTS)", level=1)
bullet("   2. Listens to the answer (via Deepgram STT)", level=1)
bullet("   3. Decides the next question based on the answer", level=1)
bullet("   4. Repeats until all questions are done", level=1)
bullet("   5. Saves the transcript when the interview ends", level=1)

h2("The Call Flow Step by Step")
para("1. User clicks 'Launch Live Interview' in Tab 5 of the app")
para("2. launcher.py starts all 3 processes above")
para("3. The app shows a URL (http://localhost:18765/livekit)")
para("4. User opens that URL in a browser")
para("5. The webpage connects to the LiveKit server using WebRTC")
para("6. The webpage publishes (sends) the user's webcam + microphone")
para("7. The AI agent also connects to the same room")
para("8. They can now see and hear each other — the interview begins")
para("9. When done, the transcript is saved and can be imported into the report")

h2("Important Technical Details")
bullet("The LiveKit server is SELF-HOSTED — it runs on your own machine, not in the cloud")
bullet("This means no monthly fees and full control, but you need to download the server binary")
bullet("The binary is downloaded from: https://github.com/livekit/livekit/releases")
bullet("It's stored in %TEMP%/livekit/livekit-server.exe")
bullet("Audio goes through: Mic → LiveKit server → AI agent → Deepgram (STT) → Gemini (LLM) → ElevenLabs (TTS) → back to you")
bullet("The whole pipeline runs in real-time with < 2 second latency")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# DETAILED LIVEKIT FILES
# ════════════════════════════════════════════════════════════
h1("19. Voice Helpers — voice.py")

bold_para("File", "core/livekit/voice.py")
bold_para("Purpose", "Provides two simple server-side functions for voice processing.")
bold_para("How it works", "Two main functions:")

para("transcribe_audio(audio_bytes):")
bullet("Takes raw audio data (WAV format)")
bullet("Sends it to Deepgram's API (model: nova-3) via HTTP POST")
bullet("Deepgram returns the transcribed text")
bullet("Example: you speak 'Hello' → audio bytes → function returns 'Hello'")

para("synthesize_speech(text):")
bullet("Takes text string")
bullet("Sends it to ElevenLabs API (model: eleven_turbo_v2_5, voice: Rachel)")
bullet("ElevenLabs returns audio data (MP3 format)")
bullet("Example: text 'Welcome to the interview' → function returns audio bytes")

para(
    "These functions are used by the web server (whisper_server.py) to handle on-demand "
    "transcription and TTS requests. They use REST APIs (not WebSocket), which is simpler."
)

# ════════════════════════════════════════════════════════════
h1("20. Voice Web Server — whisper_server.py")

bold_para("File", "core/livekit/whisper_server.py")
bold_para("Purpose", "A mini HTTP server that acts as the bridge between the browser and the AI.")
bold_para("How it works", "A multi-threaded Python HTTP server on port 18765. It handles:")

bullet("GET /token — creates a LiveKit room and generates a join token for the browser client")
bullet("GET /livekit — serves the client.html webpage (the interview UI)")
bullet("POST /transcribe — takes uploaded audio, calls voice.py's transcribe_audio(), returns text")
bullet("GET /tts — takes text as query parameter, calls voice.py's synthesize_speech(), returns audio")
bullet("POST /save_transcript — saves the interview transcript to session state")
bullet("GET /health — returns 'OK' for monitoring")

para(
    "When someone requests a token (GET /token), the server: "
    "1. Creates a new LiveKit room with a random name, "
    "2. Creates an agent dispatch (tells the LiveKit server to assign an AI agent to this room), "
    "3. Generates a secure JWT token for the browser to connect. "
    "The token has permissions to join the room, publish audio/video, and subscribe to the AI agent."
)

# ════════════════════════════════════════════════════════════
h1("21. Voice Launcher — launcher.py")

bold_para("File", "core/livekit/launcher.py")
bold_para("Purpose", "Connects the LiveKit system to the Streamlit app. Called when you click buttons in Tab 5.")
bold_para("How it works", "Three main functions:")

bullet("launch() — called when app starts. Starts the LiveKit server + web server. "
       "Returns the URL for the LiveKit client page.")
bullet("start_agent(resume, jd, room) — starts the AI voice agent as a separate process. "
       "Passes the resume and job description as environment variables so the agent knows "
       "what to ask about. Checks if agent is already running to avoid duplicates.")
bullet("cleanup() — called on app shutdown. Kills all running processes "
       "(LiveKit server, web server, agent) to avoid leftover processes.")

para(
    "The launcher manages the process lifecycle. It uses Python's subprocess module "
    "to start/stop each component. When the Streamlit app restarts (which happens often "
    "during development), cleanup() ensures no orphan processes are left running."
)

# ════════════════════════════════════════════════════════════
h1("22. Voice Agent — run_agent.py")

bold_para("File", "core/livekit/run_agent.py")
bold_para("Purpose", "The AI voice agent that conducts the interview inside a LiveKit room.")
bold_para("How it works", "This is the core of the voice system. It:")

para("1. Connects to LiveKit server and registers itself as an available agent")
para("2. When a job arrives (someone creates a room), it:")
bullet("   a. Creates an AgentSession with Deepgram STT + Gemini LLM + ElevenLabs TTS", level=1)
bullet("   b. Connects to the room where the candidate is waiting", level=1)
bullet("   c. The on_enter() method fires automatically → says greeting + first question", level=1)
bullet("   d. Starts listening to the candidate's microphone", level=1)

para("3. During the conversation:")
bullet("   • User speaks → Deepgram STT converts to text", level=1)
bullet("   • Text goes to Gemini LLM → generates response", level=1)
bullet("   • Response goes to ElevenLabs TTS → converts to speech", level=1)
bullet("   • Speech plays back to the candidate", level=1)
bullet("   • The on_user_turn_completed() method fires after each answer", level=1)

para("4. When the interview ends:")
bullet("   • The transcript (all questions + answers) is sent via data channel to the client", level=1)
bullet("   • The session closes", level=1)

para("Key components in this file:")
bullet("InterviewAgent class — custom agent with on_enter() and on_user_turn_completed()")
bullet("_build_instructions() — creates the system prompt for Gemini (includes resume + JD context)")
bullet("entrypoint(ctx) — main handler that sets up STT/LLM/TTS and starts the session")
bullet("main() — starts the AgentServer that listens for LiveKit job dispatches")

para(
    "The agent uses LiveKit's AgentSession API, which handles all the complexity of "
    "real-time audio streaming. We just need to configure the STT (Deepgram), LLM (Gemini), "
    "and TTS (ElevenLabs), and the framework takes care of the rest."
)

# ════════════════════════════════════════════════════════════
h1("23. Browser Client — client.html")

bold_para("File", "core/livekit/client.html")
bold_para("Purpose", "A web page that the candidate opens in their browser to join the video interview.")
bold_para("How it works", "This is a complete HTML page with embedded JavaScript that:")

bullet("1. Connects to the LiveKit room using LiveKit Client SDK (loaded from CDN)")
bullet("2. Requests webcam + microphone permissions")
bullet("3. Shows the candidate's webcam feed on screen")
bullet("4. Displays a transcript area showing what the AI says and what the candidate says")
bullet("5. Sends/receives data channel messages for questions, transcripts, and status updates")
bullet("6. When the interview ends, saves the transcript by sending a POST to /save_transcript")

para(
    "The page is styled with a dark theme. It automatically connects when loaded "
    "(the room token is fetched from /token with a query parameter for the room name)."
)

# ════════════════════════════════════════════════════════════
h1("24. LiveKit Server Config — livekit.yaml")

bold_para("File", "core/livekit/livekit.yaml")
bold_para("Purpose", "Configuration file for the LiveKit server binary.")
bold_para("How it works", "Tells the LiveKit server:")

bullet("port: 7880 — which port to listen on for WebSocket connections")
bullet("bind_addresses: ['127.0.0.1'] — only accept connections from localhost (secure)")
bullet("api_key: 'devkey' and api_secret: 'secret' — simple auth for development")
bullet("rtc.port_range: 7882-7892 — which ports to use for media streams")
bullet("log_level: 'info' — how much logging to show")
bullet("redis — disabled (not needed for single-machine setup)")

# ════════════════════════════════════════════════════════════
h1("25. Extra Files (Dead Code)")

para(
    "These files exist in the repository but are NOT used by the current application. "
    "They were written during earlier development phases and kept for reference."
)

h2("frontend/ folder (all files)")
para("These were an attempt to build a multi-page Streamlit app. The current app.py has "
     "everything inline, so frontend/ is abandoned. Files: components.py, pages/input_page.py, "
     "pages/analysis_page.py, pages/graph_page.py")

h2("core/gemini.py")
para("Old Gemini-only client. Superseded by core/llm.py which supports both Gemini and Ollama.")

h2("core/livekit/adapter.py")
para("An attempt to create a minimal LiveKit adapter. Never finished, never used.")

h2("core/livekit/agent_standalone.py")
para("Full standalone voice agent using faster-whisper and edge-tts (local, no cloud APIs). "
     "Superseded by run_agent.py which uses Deepgram + ElevenLabs (better quality).")

h2("core/livekit/start_livekit.py")
para("An earlier version of the launcher. Now replaced by launcher.py which integrates with Streamlit.")

h2("core/check_livekit.py, core/test_apis.py, core/test_dispatch.py")
para("Diagnostic/test scripts. Not part of the main application.")

# ════════════════════════════════════════════════════════════
h1("26. API Keys & Security")

para("The project uses 3 cloud APIs. All keys are stored in .env (gitignored, never committed):")

h2("Gemini API Key")
bullet("Used for: reading CV, reading JD, generating questions, evaluating answers")
bullet("Service: Google Gemini AI (generativelanguage.googleapis.com)")
bullet("Cost: Free tier available (60 requests/minute)")

h2("Deepgram API Key")
bullet("Used for: converting speech to text (STT)")
bullet("Service: Deepgram (api.deepgram.com)")
bullet("Model: nova-3 (most accurate)")
bullet("Cost: $0.0043 per minute (free credit for testing)")

h2("ElevenLabs API Key")
bullet("Used for: converting text to speech (TTS)")
bullet("Service: ElevenLabs (api.elevenlabs.io)")
bullet("Model: eleven_turbo_v2_5 (fastest)")
bullet("Voice: JBFqnCBsd6RMkjVDRZzb (Rachel)")
bullet("Cost: $5/month starter plan")

h2("LiveKit (self-hosted, no API key needed)")
bullet("LiveKit server runs locally — no cloud fees")
bullet("The 'devkey' and 'secret' in the config are local-only, not real secrets")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
h1("27. Data Flow Summary")

para("Here is how data flows through the system from start to finish:")

h2("Offline Phase (before interview)")
para("1. User uploads CV + JD → app.py Tab 1")
para("2. cv_agent.py extracts skills from CV using Gemini")
para("3. jd_agent.py extracts requirements from JD using Gemini")
para("4. skill_graph.py builds skill map and finds gaps")
para("5. visualize.py draws comparison graphs (Tab 3)")
para("6. question_agent.py generates question set (Tab 4)")

h2("Live Interview Phase (Tab 5)")
para("7. interview_loop.py starts — picks first skill using traversal.py")
para("8. question_agent.py generates a question for that skill")
para("9. Candidate answers (text/voice/video)")
para("10. evaluator.py scores the answer using Gemini")
para("11. State is updated in state.py")
para("12. If incomplete → loop back to step 7 (next skill)")
para("13. If complete → proceed to report")

h2("Report Phase (Tab 6)")
para("14. generator.py takes all scores + answers")
para("15. Generates per-skill breakdown, overall score, verdict")
para("16. Displays in Tab 6 with hiring recommendation")

para("")
para("— End of Guide —")

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "PROJECT_GUIDE.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
