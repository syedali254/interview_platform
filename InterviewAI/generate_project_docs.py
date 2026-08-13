"""Generate PROJECT_DOCS.docx — the handover document for InterviewAI.

Run: python generate_project_docs.py

Everything in here is written so that someone who has never seen the project
can run it, understand what each module does, and — most importantly — trace
every number on the final report back to the code that produced it.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.35


# ── Helpers ──────────────────────────────────────────────────────────────

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)


def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = 'Arial'
    r.bold = bold
    r.italic = italic
    p.paragraph_format.line_spacing = 1.35
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.name = 'Arial'
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = 'Arial'


def code(text):
    """Monospaced block, used for formulas and commands."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1a, 0x35, 0x5e)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = 'Arial'
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ═════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('InterviewAI')
r.bold = True
r.font.size = Pt(30)
r.font.name = 'Arial'

doc.add_paragraph()
tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run('Project Documentation and Handover Guide')
r2.font.size = Pt(17)
r2.font.name = 'Arial'

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run(
    'An Intelligent Multi-Agent AI Interview Platform\n'
    'Voice Interviewing  |  ESCO Skill Graph  |  LLM Answer Evaluation\n'
    'Attention, Posture and Vocal Analysis  |  Behavioural Integrity Detection'
)
r3.font.size = Pt(11.5)
r3.font.name = 'Arial'
r3.italic = True

doc.add_paragraph()
tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tp4.add_run('CMP7200 — MSc Computer Science Individual Masters Project\n'
                 'Birmingham City University')
r4.font.size = Pt(10.5)
r4.font.name = 'Arial'

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════
h1('1. What This Project Is')
# ═════════════════════════════════════════════════════════════════════════

para(
    'InterviewAI is a working platform that conducts a technical job interview by voice, '
    'scores the candidate\'s answers, monitors how they behaved during the session, and '
    'produces a single evidence-backed hiring recommendation. It is built as a research '
    'artefact for an MSc dissertation, not as a commercial hiring product, and it is not '
    'intended to make real employment decisions.'
)

para(
    'The problem it addresses is that commercial AI interview tools return scores without '
    'explanation. Candidates are ranked by systems they cannot interrogate, and the EU AI '
    'Act now classifies exactly this class of system as high-risk, requiring transparency '
    'and human oversight. Every number this system produces is therefore traceable: the '
    'report shows the rubric breakdown behind each answer score, the features behind the '
    'integrity verdict, and the arithmetic behind the final recommendation.'
)

h2('1.1 What happens in one sentence')
para(
    'The candidate uploads a CV and a job description; the system maps both onto a '
    'standardised skill taxonomy to find the gaps; it generates interview questions '
    'targeting those gaps; an AI interviewer conducts a live voice interview; every answer '
    'is scored against a generated reference answer; attention, posture, voice and session '
    'behaviour are measured throughout; and all of it is fused into a final report.'
)

h2('1.2 Who this document is for')
para(
    'Anyone taking over this codebase. Section 2 gets the system running. Sections 3 and 4 '
    'explain the architecture. Section 5 is the important one — it defines exactly how every '
    'score and percentage on the final report is calculated, with the constants and the file '
    'each formula lives in. Sections 6 to 9 cover the code layout, the API, configuration, '
    'and what is deliberately not implemented.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('2. Getting It Running')
# ═════════════════════════════════════════════════════════════════════════

h2('2.1 Prerequisites')
table(
    ['Requirement', 'Where to get it', 'Notes'],
    [
        ['Python 3.11+', 'Installed automatically', 'run.bat installs it via winget if missing'],
        ['Node.js LTS', 'Installed automatically', 'run.bat installs it via winget if missing'],
        ['GEMINI_API_KEY', 'aistudio.google.com/apikey', 'All language model work'],
        ['DEEPGRAM_API_KEY', 'console.deepgram.com', 'Speech to text, and the fallback voice'],
        ['ELEVENLABS_API_KEY', 'elevenlabs.io', 'OPTIONAL - preferred voice. Falls back to Deepgram if absent'],
    ],
    widths=[4.0, 5.5, 7.0],
)

h2('2.2 Running it')
para('From the project root folder, double-click run.bat, or from a terminal:')
code('run.bat')
para(
    'If Python or Node.js are missing, run.bat installs them with winget, which ships with '
    'Windows 10 1709 and later. After an automatic install the script asks you to close the '
    'window and run it once more, so Windows picks up the new PATH. If winget is not '
    'available it says exactly what to install by hand.'
)
para(
    'The first run creates the virtual environment, installs Python and Node packages, '
    'downloads the MediaPipe vision models, prompts for the three API keys, and builds the '
    'web app. Later runs just start the server. Then open http://localhost:8000 in Chrome '
    'or Edge and allow camera and microphone access.'
)

h2('2.3 Manual setup, if run.bat fails')
code(
    'cd InterviewAI\n'
    'python -m venv venv\n'
    'venv\\Scripts\\activate\n'
    'pip install -r requirements.txt\n'
    'cd frontend\n'
    'npm install\n'
    'npm run build          (this also downloads the MediaPipe models)\n'
    'cd ..\n'
    'copy .env.example .env  (then edit .env and paste your three API keys)\n'
    'python server.py'
)

h2('2.4 What runs where')
para(
    'One Python process serves both the API and the built web app on port 8000. When an '
    'interview starts, the server downloads and launches a local LiveKit media server on '
    'port 7880, and spawns the voice agent as a separate Python process. Both are shut down '
    'when the interview ends. Nothing needs to be installed manually for this — the launcher '
    'fetches the LiveKit binary on first use.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('3. End-to-End Flow')
# ═════════════════════════════════════════════════════════════════════════

para('The user moves through six screens. Each one calls the API and stores its result in the session.')

table(
    ['Step', 'Screen', 'What happens', 'Modules'],
    [
        ['1', 'Upload & Parse', 'CV (PDF or text) and job description are parsed into structured data', 'M1, M2'],
        ['2', 'Skill Graph', 'Both skill sets are mapped onto the ESCO taxonomy and compared', 'M3'],
        ['3', 'Questions', 'Questions are generated for the gaps and ordered by priority', 'M4'],
        ['4', 'Setup & Mode', 'Candidate picks voice or text, tests devices, and acknowledges what is monitored', '—'],
        ['5', 'Live Interview', 'Voice or text interview runs; attention, posture and behaviour are recorded throughout', 'M5, M7, M8, M10'],
        ['6', 'Report', 'Every answer is scored, behaviour assessed, and everything fused', 'M6, M9, M11, M12'],
    ],
    widths=[1.3, 3.2, 8.5, 3.5],
)

h2('3.1 What happens during the live interview')
para(
    'The agent process joins a LiveKit room alongside the candidate\'s browser. Deepgram '
    'transcribes the candidate continuously. Gemini generates the interviewer\'s replies. '
    'ElevenLabs speaks them.'
)
para(
    'Each interviewer utterance is buffered in full before anything is spoken. The complete '
    'text is published to the browser first, the browser renders it, and only then does '
    'audio synthesis begin. This ordering is deliberate: the candidate reads the question '
    'before hearing it. It also fixed a real defect — streaming partial language-model '
    'tokens straight into the text-to-speech socket caused the agent to speak only the '
    'first few words of each reply while the full text still appeared on screen.'
)
para(
    'While this happens, the browser runs face and pose detection five times a second and '
    'analyses microphone audio ten times a second for pitch and energy, folding both into '
    'one sample per second for the report. Tab switches are recorded separately. All of it '
    'is sent to the agent over the data channel and also kept in the browser.'
)
para(
    'What is detected is drawn on screen. The candidate sees their own face outline, the eye '
    'points the gaze estimate uses, a box around their face carrying a live attention '
    'percentage, and the skeleton of their upper body behind the posture score. The overlay '
    'can be switched off from the control bar. Showing it is deliberate: a system that '
    'measures someone should let them see what it is measuring.'
)
para(
    'The conversation is shown beside the video as it happens. Each interviewer question and '
    'each transcribed answer appears as its own message, and whatever the candidate is saying '
    'right now appears underneath in grey until it is finalised.'
)

h3('Why the interview starts almost immediately')
para(
    'Pressing Begin Interview used to mean waiting more than twenty seconds. Measuring it '
    'showed the delay was not the network. It was Python: loading the livekit-agents library '
    'and its plugins takes about twelve seconds before the agent process can do anything.'
)
para(
    'The device-setup screen now calls /api/prewarm the moment it opens. That starts the '
    'media server, starts the agent process, and lets it connect to a room and check its '
    'voice provider, all while the candidate is reading the briefing and testing their '
    'camera. The agent then sits in the room and waits. When the candidate finally presses '
    'Begin, the only work left is issuing a token and joining, which takes about one second. '
    'A three-two-one countdown covers it.'
)
para(
    'If prewarm never ran, or the agent process died in the meantime, the token endpoint '
    'notices and starts a fresh agent instead, so the interview still works. It just takes '
    'the old amount of time.'
)

h2('3.2 Two interview modes')

para(
    'The candidate chooses between a voice interview and a text interview on the setup '
    'screen. They are the same interview. Both use one shared set of interviewer '
    'instructions (core/agents/interviewer_prompt.py), the same question bank in the same '
    'graph-priority order, the same question and time budgets, and the same ending rules, '
    'and both produce the same transcript. The report pipeline never learns which one ran.'
)

table(
    ['', 'Voice mode', 'Text mode'],
    [
        ['How the question arrives', 'Displayed, then spoken aloud', 'Written into the chat'],
        ['How the candidate answers', 'Speaks; Deepgram transcribes', 'Types; Enter sends'],
        ['Transport', 'LiveKit room and agent subprocess', 'One HTTP request per answer'],
        ['Devices needed', 'Camera and microphone', 'Camera only'],
        ['M7 attention', 'Yes', 'Yes'],
        ['M8 posture', 'Yes', 'Yes'],
        ['M10 vocal delivery', 'Yes', 'Not applicable - nothing is spoken'],
        ['Extra integrity signal', 'Speech hesitation', 'Pasted answers are flagged'],
        ['Scoring and report', 'Identical', 'Identical'],
    ],
    widths=[4.4, 6.0, 6.1],
)

para(
    'Text mode has no vocal delivery to measure, so that signal is absent rather than zero. '
    'The fusion engine renormalises the engagement weights across whichever presence signals '
    'actually arrived, so a typed interview still yields a complete engagement score from '
    'attention and posture instead of one with a hole in it. In place of vocal analysis, '
    'text mode contributes an integrity signal of its own: an answer that was pasted rather '
    'than typed is flagged and shown on the report.'
)

para(
    'Only voice mode needs the media server and the agent subprocess, so only voice mode '
    'triggers the prewarm described above. A text interview begins as soon as the first '
    'request returns.'
)

h2('3.3 How the interview ends')
para('There are three ways an interview finishes, and all of them end the same way.')
bullet('the agent gives a short closing statement, then the report screen appears automatically.',
       bold_prefix='All questions covered — ')
bullet('the agent asks one confirmation question. If the candidate confirms, it thanks them and ends. '
       'If they decline, the interview continues and it does not ask again.',
       bold_prefix='Candidate asks to stop — ')
bullet('once MAX_INTERVIEW_QUESTIONS or INTERVIEW_TIME_BUDGET_MINS is reached, a watchdog '
       'instructs the agent to wrap up. If the model fails to close within 35 seconds, the '
       'session is closed anyway.',
       bold_prefix='Budget reached — ')
para(
    'In every case the agent calls its end_interview tool, the browser shows a "Thank you '
    'for your time" card, the transcript is saved, the LiveKit server and agent process are '
    'shut down, and the report screen loads and begins scoring.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('4. Module Map')
# ═════════════════════════════════════════════════════════════════════════

para('Twelve modules across four phases, as defined in the project proposal.')

table(
    ['#', 'Module', 'Technology', 'Where it lives', 'Status'],
    [
        ['M1', 'CV parsing', 'Gemini + PyMuPDF', 'core/agents/cv_agent.py', 'Implemented'],
        ['M2', 'Job description parsing', 'Gemini', 'core/agents/jd_agent.py', 'Implemented'],
        ['M3', 'Skill knowledge graph', 'NetworkX + ESCO', 'core/graph/skill_graph.py', 'Implemented'],
        ['M4', 'Question generation', 'Gemini + graph priority', 'core/agents/question_agent.py', 'Implemented'],
        ['M5', 'Voice interview', 'LiveKit, Deepgram, ElevenLabs', 'core/livekit/run_agent.py', 'Implemented'],
        ['M6', 'Answer evaluation', 'Gemini LLM-as-Judge', 'core/evaluator/evaluator.py', 'Implemented'],
        ['M7', 'Attention monitoring', 'MediaPipe FaceLandmarker', 'frontend/src/lib/vision.js', 'Implemented'],
        ['M8', 'Posture analysis', 'MediaPipe PoseLandmarker', 'frontend/src/lib/vision.js', 'Implemented'],
        ['M9', 'Behavioural integrity', 'Isolation Forest', 'core/evaluator/integrity.py', 'Implemented'],
        ['M10', 'Vocal delivery analysis', 'Web Audio prosody', 'frontend/src/lib/voice.js', 'Implemented'],
        ['M11', 'Weighted fusion', 'Rules + weights', 'core/evaluator/fusion.py', 'Implemented'],
        ['M12', 'Report assembly', 'Structured templates', 'core/report/generator.py', 'Implemented'],
    ],
    widths=[1.0, 3.6, 4.2, 5.0, 2.7],
)

h2('4.1 A note on M6, M7, M8 and M10')
para(
    'The proposal originally specified a second answer-scoring track — Sentence-BERT '
    'embeddings feeding an XGBoost classifier with SHAP explanations — to be compared '
    'against the language model judge. That track was built, trained, measured and then '
    'removed. Two findings drove the decision. First, the comparison was circular: the '
    'classifier\'s training labels were themselves generated by a language model, so '
    'agreement between the two tracks was guaranteed by the experimental design rather '
    'than earned. Second, the trained model failed on inspection — an answer identical to '
    'the reference scored 64.7 out of 100, and a correct paraphrase of it scored 39.2, '
    'below the threshold at which the system reports a skill gap. The full evidence and '
    'the revised objectives are recorded in docs/track-b-rejection.md. Answer evaluation '
    'is done entirely by the Gemini judge described in Section 5.1.'
)
para(
    'M7 and M8 use MediaPipe as specified. M10 was specified as a wav2vec2 speech-emotion '
    'classifier; it is implemented instead as prosodic analysis — pitch, energy, pause and '
    'fluency measurement — computed locally in the browser. This substitution is deliberate: '
    'it requires no model download, runs offline, keeps all audio on the candidate\'s '
    'machine, and every component of the resulting score is inspectable, which suits an '
    'explainability-focused system better than an opaque emotion label. Facial emotion is '
    'additionally captured via face-api.js and reported as context only — it does not feed '
    'the score.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('5. Scoring — How Every Number Is Calculated')
# ═════════════════════════════════════════════════════════════════════════

para(
    'This section is the reference for the whole assessment. Each subsection gives the '
    'formula, the constants, and the file the code is in. All scores are on a 0-100 scale '
    'unless stated otherwise.',
    bold=True,
)

# ── 5.1 ──────────────────────────────────────────────────────────────────
h2('5.1 Answer Score (M6) — core/evaluator/evaluator.py')

h3('Step 1 — generate a reference answer')
para(
    'Before the candidate\'s answer is looked at, Gemini is asked to write an ideal answer '
    'to the same question, capped at 100 words and phrased the way a strong candidate would '
    'speak it. This reference is what the candidate is measured against. It is regenerated '
    'per question, so the standard adapts to the question\'s difficulty rather than being '
    'fixed in advance.'
)

h3('Step 2 — score against a four-criterion rubric')
para('The judge scores four criteria, each out of 25, giving a natural 0-100 total.')

table(
    ['Criterion', 'Out of', 'What it measures'],
    [
        ['technical_accuracy', '25', 'Is everything the candidate actually said correct? Only wrong or misleading statements lose marks. Omissions do not.'],
        ['completeness', '25', 'Did they cover the essential points in the reference? Only genuinely missing concepts lose marks. Brevity does not.'],
        ['clarity', '25', 'Is the explanation well structured and easy to follow?'],
        ['relevance', '25', 'Does it address the question that was actually asked?'],
    ],
    widths=[4.0, 1.8, 10.7],
)

para(
    'Accuracy and completeness are deliberately independent. An earlier version forced '
    'accuracy below 15/25 whenever any reference point was missing, which meant a short but '
    'entirely correct answer was penalised twice for the same thing and scored 8/25 on '
    'accuracy. Separating them means the report can distinguish a candidate who is right but '
    'brief from one who is confidently wrong — these are very different signals to a '
    'recruiter.'
)
para(
    'The rubric also states that clarity and relevance cannot be high when the answer '
    'contains no substance, since a candidate cannot clearly explain nothing. Without that '
    'rule a content-free answer scored around 37/100 purely for being fluent and on-topic; '
    'with it, such answers score around 11.'
)

h3('Step 3 — score twice, in two different criterion orders')
para(
    'Language-model judges are known to over-weight whichever criterion they read first '
    '(Stureborg et al., 2024). Each answer is therefore scored twice with the four criteria '
    'presented in two different orders, and the two totals are averaged.'
)
code(
    'order 1 = technical_accuracy, completeness, clarity, relevance\n'
    'order 2 = clarity, relevance, technical_accuracy, completeness\n\n'
    'final_score = (total_call_1 + total_call_2) / 2\n'
    'spread      = | total_call_1 - total_call_2 |'
)

h3('Step 4 — turn the spread into a reliability signal')
para(
    'The gap between the two calls is how much the judge disagreed with itself on that '
    'answer. It is reported per answer and aggregated for the session.'
)
table(
    ['Spread', 'Consistency', 'Meaning'],
    [
        ['< 8 points', 'high', 'Stable judgement; the score can be relied on'],
        ['8 to 16 points', 'moderate', 'Some instability; treat with mild caution'],
        ['16 points or more', 'low', 'Judge was unstable — the answer is flagged for human review'],
    ],
    widths=[3.8, 3.0, 9.7],
)

h3('Step 5 — verdict band')
code(
    'score >= 70  ->  "strong"     (SCORE_STRONG_THRESHOLD, core/config.py)\n'
    'score >= 40  ->  "weak"       (SCORE_WEAK_THRESHOLD,   core/config.py)\n'
    'otherwise    ->  "gap"'
)

h3('Which answers get scored')
para(
    'Every substantive answer is scored. Each exchange is first classified as technical, '
    'behavioural, or logistics. Logistics turns — greetings, "are you ready", sign-offs — are '
    'not scored, because scoring "Yes, I\'m ready" as a failed technical answer would '
    'corrupt the average. Everything classified technical or behavioural is scored regardless '
    'of length, so a one-line non-answer appears in the report with a low score rather than '
    'being silently dropped. Only replies shorter than three words are treated as empty. '
    'Classification is done in one language-model call for the whole transcript, with a '
    'keyword fallback if that call fails.'
)

h3('The overall answer score')
code('overall_score = mean(final_score of every scored answer)')

# ── 5.2 ──────────────────────────────────────────────────────────────────
h2('5.2 Skill Match Percentage (M3) — core/graph/skill_graph.py')

para(
    'The percentage shown on the Skill Graph screen and used as the skill_coverage component '
    'of the final score is simply how many of the role\'s required skills were found on the CV:'
)
code(
    'match_percentage = |candidate_skills  INTERSECT  required_skills|\n'
    '                   ------------------------------------------------  x 100\n'
    '                              |required_skills|'
)
para(
    'Both sets are sets of taxonomy node identifiers, not raw strings, so "K8s" on the CV and '
    '"Kubernetes" in the job description count as the same skill. Nice-to-have skills are '
    'tracked separately and never affect this percentage.'
)

h3('How a skill string becomes a node')
para(
    'Matching is deliberately conservative. A CV skill is mapped to a taxonomy concept only '
    'if one of these succeeds, in order:'
)
bullet('the normalised text equals a preferred label exactly', bold_prefix='Exact label — ')
bullet('it equals a known alias or abbreviation ("k8s", "postgres", "reactjs")', bold_prefix='Alias — ')
bullet('it equals a label with its parenthetical stripped, so "Python" reaches '
       '"Python (computer programming)"', bold_prefix='Base form — ')
bullet('close string similarity, but only for text of 6 characters or more, at a 0.88 cutoff',
       bold_prefix='Fuzzy — ')
para(
    'If none succeed the skill becomes its own node rather than being forced onto an '
    'unrelated concept. An earlier version had a bare substring fallback, which mapped '
    '"Team Leadership", "Leadership", "Scrum" and "Problem Solving" all onto the ESCO skill '
    '"R", and "Communication" onto "telecommunications engineering". That fallback has been '
    'removed. Short labels such as "R", "C#" and "SQL" must now match exactly or not at all, '
    'which is why fuzzy matching has a minimum length.'
)
para(
    'Unmatched skills from the CV and the job description share one namespace, so a skill the '
    'taxonomy does not know about still registers as matched when it appears on both sides.'
)

h3('The five node statuses')
table(
    ['Status', 'Meaning', 'Colour on screen'],
    [
        ['matched', 'On the CV and required by the role', 'Green'],
        ['missing', 'Required by the role, absent from the CV — a gap', 'Red'],
        ['bonus', 'A nice-to-have the candidate already has', 'Violet'],
        ['bonus_missing', 'A nice-to-have the candidate lacks', 'Amber'],
        ['extra', 'On the CV but not asked for by the role', 'Slate'],
    ],
    widths=[3.2, 9.3, 4.0],
)
para(
    'The "Skills in play" figure on the report is the count of distinct nodes across all five '
    'statuses. It is not candidate skills plus required skills, because a matched skill '
    'belongs to both sides and must not be counted twice.'
)

h3('Interview topic priority')
para(
    'Topics are drawn from the gap analysis: up to three missing required skills as high '
    'priority, up to three matched required skills as medium (to verify depth of knowledge), '
    'and up to two matched nice-to-haves as low. Question generation then orders technical '
    'questions by this priority so genuine gaps are probed before the time budget runs out.'
)

# ── 5.3 ──────────────────────────────────────────────────────────────────
h2('5.3 Behavioural Integrity (M9) — core/evaluator/integrity.py')

para(
    'An Isolation Forest — an unsupervised anomaly detector — is trained on synthetic normal '
    'interview behaviour, then asked how unusual the real session looks. It never decides '
    'anything on its own; it annotates the session for the recruiter.'
)

h3('The eight features')
table(
    ['Feature', 'How it is measured', 'Normal baseline'],
    [
        ['avg_response_time_sec', 'Mean seconds from question start to answer end', 'N(32, 11)'],
        ['response_time_std', 'Variation in those times', 'N(14, 6)'],
        ['tab_switches', 'Times the candidate left the browser tab', 'Poisson(0.7)'],
        ['inactivity_ratio', 'Sustained disengaged time / session length', 'N(0.04, 0.03)'],
        ['avg_answer_length_words', 'Mean words per answer', 'N(48, 18)'],
        ['answer_length_variance_coeff', 'Std / mean of answer lengths', 'N(0.45, 0.18)'],
        ['hesitation_ratio', 'Filler words per word of answer', 'N(0.04, 0.03)'],
        ['engagement_score', 'Derived from tab switches and inactivity', 'U(0.6, 1.0)'],
    ],
    widths=[5.2, 7.3, 4.0],
)
para(
    'The baseline distributions must match what the system actually measures. Response time '
    'here is measured from the interviewer\'s question starting to the candidate\'s answer '
    'finishing, so it includes both the question being spoken and the answer being delivered '
    '— typically 20 to 60 seconds. An earlier baseline assumed a 5 to 15 second "thinking '
    'time" reading and consequently flagged every ordinary session as anomalous, reporting '
    '"flagged" with no risk factors to explain it.'
)
para(
    'Natural gaze behaviour is also not treated as inactivity. Only the proportion of the '
    'session spent looking away above 20 percent counts as idle time, because people '
    'routinely break eye contact while thinking.'
)

h3('Model and calibration')
code(
    'IsolationForest(n_estimators=200, contamination=0.05, random_state=42)\n'
    'fitted on 400 synthetic normal sessions\n\n'
    'raw = model.decision_function(session_features)   # roughly -0.3 .. +0.2\n\n'
    'p01, p99 = 1st and 99th percentile of raw scores across the baseline\n\n'
    'integrity_score = clamp( 50 + (raw - p01) / (p99 - p01) * 50 , 0 , 100 )'
)
para(
    'The raw decision function is not a number a person can interpret, so it is calibrated '
    'against the baseline distribution: the 1st percentile of normal sessions maps to 50 and '
    'the 99th to 100, with genuine anomalies falling below 50. The calibration values are '
    'stored inside the saved model file, and a model saved before calibration existed is '
    'rejected and retrained automatically.'
)

h3('Verdict bands')
code(
    'integrity_score >= 60  ->  "normal"\n'
    'integrity_score >= 35  ->  "suspicious"\n'
    'otherwise              ->  "flagged"'
)

h3('Risk factors')
para('These are the specific behaviours named on the report. They are thresholds, evaluated independently of the model:')
table(
    ['Condition', 'Reported as'],
    [
        ['avg_response_time_sec < 6', 'Implausibly fast responses (possible pre-written answers)'],
        ['avg_response_time_sec > 90', 'Very long response times (possible external assistance)'],
        ['tab_switches > 5', 'High tab-switch count'],
        ['inactivity_ratio > 0.3', 'Extended inactivity during the session'],
        ['answer_length_variance_coeff > 1.2', 'Highly inconsistent answer lengths (possible mixed sources)'],
        ['engagement_score < 0.4', 'Low engagement throughout the session'],
    ],
    widths=[6.5, 10.0],
)
para(
    'If the verdict is adverse but no threshold was crossed, the report says so explicitly '
    'rather than showing a bare "flagged" with nothing behind it. A verdict a recruiter '
    'cannot act on is worse than no verdict.'
)

# ── 5.4 ──────────────────────────────────────────────────────────────────
h2('5.4 Presence: Attention, Posture and Voice (M7, M8, M10)')

para(
    'All three run entirely in the candidate\'s browser. No video or audio is recorded, '
    'stored or transmitted — only the derived per-second numbers leave the machine.'
)

h3('M7 Attention — frontend/src/lib/vision.js')
para(
    'MediaPipe FaceLandmarker provides 478 face landmarks per frame, sampled once per second. '
    'Head orientation is derived from landmark geometry rather than the raw transformation '
    'matrix, because the ratios are stable across different cameras and are easier to justify:'
)
code(
    'eye_mid   = midpoint of the two outer eye corners (landmarks 33 and 263)\n'
    'eye_dist  = distance between them\n'
    'nose      = landmark 1\n\n'
    'yaw   = (nose.x - eye_mid.x) / eye_dist        horizontal head turn\n'
    'pitch = (nose.y - eye_mid.y) / eye_dist        vertical head tilt'
)
para(
    'The first thirty detections are used to calibrate the candidate\'s own neutral pose, so '
    'the score measures deviation from how they naturally sit rather than from an assumed '
    'ideal. Detection runs at five hertz so the on-screen overlay tracks smoothly; every '
    'fifth detection is averaged into the one-per-second sample the report uses.'
)
code(
    'yaw_dev   = |yaw   - yaw_baseline|   / 0.42\n'
    'pitch_dev = |pitch - pitch_baseline| / 0.38\n\n'
    'attention = clamp01( 1 - (yaw_dev + pitch_dev) / 2 )        0.0 to 1.0\n\n'
    'looking_away = attention < 0.45\n'
    '4 consecutive looking-away samples raise a distraction event'
)

h3('M8 Posture — frontend/src/lib/vision.js')
para('MediaPipe PoseLandmarker provides body landmarks. Three deviations are measured, each normalised by shoulder width so distance from the camera does not matter:')
code(
    'shoulder_width = distance between left and right shoulder\n\n'
    'tilt   = |left_shoulder.y - right_shoulder.y| / shoulder_width      / 0.22\n'
    'slouch = drop in head height below the calibrated baseline          / 0.35\n'
    'lean   = |nose.x - shoulder_midpoint.x| / shoulder_width            / 0.30\n\n'
    'posture = clamp01( 1 - (tilt + slouch + lean) / 3 )'
)
para('Any component exceeding its tolerance is named on the report as "shoulders uneven", "slouching" or "leaning off-centre", along with the proportion of the session it applied to.')

h3('M10 Vocal delivery — frontend/src/lib/voice.js')
para(
    'The microphone stream is analysed ten times per second with the Web Audio API. '
    'Fundamental frequency is estimated by autocorrelation bounded to the human speech range '
    '(75-400 Hz), which keeps it cheap enough to run on the main thread. Frames quieter than '
    '0.012 RMS are treated as silence.'
)
para('Four components, each scored 0 to 1 and equally weighted:')
table(
    ['Component', 'Formula', 'What it captures'],
    [
        ['projection', 'clamp01(avg_energy / 0.03)', 'Vocal loudness; quiet delivery reads as tentative'],
        ['fluency', 'full marks for 30-75% voiced time, scaling down outside that band',
         'Proportion of the turn actually spent speaking'],
        ['expression', 'full marks for 10-45 Hz pitch variation, scaling down outside',
         'Monotone and erratic delivery both score below naturally varied speech'],
        ['composure', 'clamp01(1 - long_pauses_per_minute / 4)',
         'Hesitation; a pause over 1.2 seconds counts as long'],
    ],
    widths=[3.0, 6.5, 7.0],
)
code('vocal_confidence = (projection + fluency + expression + composure) / 4  x  100')

h3('Combining them into the engagement score')
para('The three presence modules are combined into the single engagement component used by the final fusion:')
code(
    'ENGAGEMENT_WEIGHTS = { attention: 0.45, posture: 0.20, voice: 0.35 }\n\n'
    'weighted_mean = sum(source x weight) / sum(weights of available sources)\n\n'
    'distraction_penalty = min( distraction_events x 4 , 30 )\n\n'
    'engagement = clamp( weighted_mean - distraction_penalty , 0 , 100 )'
)
para(
    'Weights are renormalised over whatever sources actually produced data, so a session '
    'where posture could not be measured still yields a valid engagement score from attention '
    'and voice alone. If none of the three produced data — an older browser, or the MediaPipe '
    'models were unavailable — the score falls back to an estimate from distraction events '
    'only, and the report marks it as estimated rather than measured.'
)

# ── 5.5 ──────────────────────────────────────────────────────────────────
h2('5.5 Final Fusion and Recommendation (M11) — core/evaluator/fusion.py')

para('Four components, weighted to sum to 1.0:')
table(
    ['Component', 'Weight', 'Source', 'Meaning'],
    [
        ['answer_quality', '0.50', 'M6', 'Mean of every scored answer'],
        ['skill_coverage', '0.20', 'M3', 'Skill match percentage'],
        ['behavioral_integrity', '0.15', 'M9', 'Calibrated integrity score'],
        ['engagement', '0.15', 'M7 + M8 + M10', 'Attention, posture and vocal delivery'],
    ],
    widths=[4.3, 2.0, 3.2, 7.0],
)
code(
    'fusion_score =   answer_quality       x 0.50\n'
    '               + skill_coverage       x 0.20\n'
    '               + behavioral_integrity x 0.15\n'
    '               + engagement           x 0.15'
)

h3('Recommendation bands')
table(
    ['Condition', 'Recommendation', 'Label', 'Confidence'],
    [
        ['integrity < 30', 'disqualified', 'Session Integrity Compromised', 'high'],
        ['fusion >= 72', 'strong_hire', 'Strong Hire', 'high if >= 80, else moderate'],
        ['fusion >= 55', 'hire', 'Hire — Meets Requirements', 'moderate'],
        ['fusion >= 40', 'consider', 'Consider — Development Needed', 'moderate'],
        ['below 40', 'no_hire', 'No Hire — Significant Gaps', 'high if < 25, else moderate'],
    ],
    widths=[3.4, 3.4, 6.2, 3.5],
)
para(
    'The integrity check is an override, evaluated before the bands. A session whose '
    'integrity score falls below 30 is reported as disqualified regardless of how good the '
    'answers were, because the answers cannot be trusted to be the candidate\'s own.'
)

h3('Worked example')
para('A real run of the system produced these numbers:')
code(
    'answer_quality        59.8  x 0.50  =  29.9\n'
    'skill_coverage        57.1  x 0.20  =  11.4\n'
    'behavioral_integrity  78.5  x 0.15  =  11.8\n'
    'engagement            68.4  x 0.15  =  10.3\n'
    '                                      ------\n'
    'fusion_score                           60.9   ->  "Hire - Meets Requirements"\n\n'
    'where engagement 68.4 came from:\n'
    '  attention 79.0 x 0.45 + posture 62.0 x 0.20 + voice 81.2 x 0.35\n'
    '    = 76.4 weighted mean\n'
    '  minus 2 distraction events x 4 = 8.0 penalty\n'
    '    = 68.4'
)
para(
    'Every one of these intermediate values is present in the API response and shown on the '
    'report screen, including the per-source engagement breakdown, so the arithmetic can be '
    'checked by hand.'
)

# ── 5.6 ──────────────────────────────────────────────────────────────────
h2('5.6 Session Reliability Statistics (M12) — core/report/generator.py')

para(
    'Because the scores come from a language model, the report includes how self-consistent '
    'that model was across the whole session:'
)
table(
    ['Statistic', 'Definition'],
    [
        ['n', 'Number of answers scored'],
        ['mean_spread', 'Average gap between the two rubric-order calls, in points'],
        ['max_spread', 'Largest such gap in the session'],
        ['consistency_distribution', 'Count of answers at high, moderate and low consistency'],
        ['flagged_for_review', 'Answers whose spread reached 16 points or more'],
    ],
    widths=[5.0, 11.5],
)
para(
    'A session where the judge frequently disagreed with itself is reported as such rather '
    'than presented as a confident assessment. In testing, an ordinary answer produces a '
    'spread of 0 to 3 points, while a deliberately misleading answer — confidently stated but '
    'factually wrong — produced a spread of 20 and was automatically flagged.'
)


# ═════════════════════════════════════════════════════════════════════════
h2('5.7 Voice Provider Selection and Failure Handling')

para(
    'Text-to-speech is the one part of the system that depends on a third-party quota, and '
    'it fails in a way that is easy to misread. A provider whose credits are exhausted still '
    'accepts the connection and completes the request — it simply returns no audio. The '
    'symptom is an interviewer that speaks normally at the start of a session and then goes '
    'silent, while the transcript keeps scrolling correctly. This is not a bug in the agent; '
    'it is the free-tier character limit running out mid-interview.'
)
para(
    'The agent therefore verifies its voice provider before the interview starts, by '
    'synthesising two characters and checking that audio frames actually come back. '
    'Providers are tried in order:'
)
table(
    ['Order', 'Provider', 'Model', 'Key'],
    [
        ['1', 'ElevenLabs', 'eleven_turbo_v2_5', 'ELEVENLABS_API_KEY'],
        ['2', 'Deepgram Aura', 'aura-2-andromeda-en', 'DEEPGRAM_API_KEY'],
        ['3', 'None — text only', 'Questions displayed, not spoken', '—'],
    ],
    widths=[2.0, 4.5, 6.0, 4.0],
)
para(
    'The chosen provider is logged to agent_debug.log and sent to the browser, which shows '
    '"text only" beside the interviewer if no provider worked. The interview still runs in '
    'that state: questions appear on screen, answers are still transcribed and scored, and '
    'the report is unaffected. Voice quality degrades; the assessment does not.'
)
para(
    'To check an ElevenLabs quota directly, POST to '
    'https://api.elevenlabs.io/v1/text-to-speech/{voice_id} with the xi-api-key header. A '
    'quota_exceeded response names the exact shortfall. Set ELEVENLABS_VOICE_ID or '
    'DEEPGRAM_TTS_MODEL in .env to change voices.'
)

h2('5.8 Reading the Report Screen')

para(
    'The report is the deliverable, so it takes the full width of the page with no wizard '
    'sidebar. Everything on it is a percentage — there are no formulas printed on screen. '
    'The maths behind each figure is in the sections above; the screen just shows results.'
)

table(
    ['Section', 'What it tells you'],
    [
        ['Recommendation banner', 'The final verdict, the answer score, the fused score and how confident the system is'],
        ['Score Breakdown', 'The four components that produced the fused score, each as a percentage'],
        ['Presence', 'Attention, posture and vocal confidence, plus how they combined into engagement'],
        ['Behavioural Integrity', 'The integrity verdict, the eight measured features, and any risk factors'],
        ['Scoring Reliability', 'How consistent the judge was with itself, and how many answers were flagged'],
        ['Skill Breakdown', 'Every skill that was probed, scored worst first'],
        ['Answer-by-Answer', 'Each question, the answer, the reference answer, and the rubric percentages'],
        ['Session metrics and charts', 'Duration, exchanges, distractions, emotion timeline and distribution'],
        ['Transcript', 'The full conversation with timestamps and response times'],
    ],
    widths=[4.6, 11.9],
)

h3('Saving the report as a PDF')
para(
    'The Download PDF button expands every collapsed answer card, waits for the page to '
    'redraw, and then opens the browser print dialog, where "Save as PDF" is the '
    'destination. A print stylesheet removes the buttons and page chrome, forces background '
    'colours to render so the bars are visible on paper, and stops cards being split across '
    'a page break. This uses the browser own PDF engine, so there is no extra dependency '
    'to install and the output matches what is on screen.'
)

h1('6. Code Layout')
# ═════════════════════════════════════════════════════════════════════════

table(
    ['Path', 'What it does'],
    [
        ['run.bat', 'One-click setup and launch from the project root'],
        ['server.py', 'FastAPI backend: all endpoints, serves the built web app'],
        ['requirements.txt', 'Python dependencies'],
        ['.env', 'The three API keys (never committed)'],
        ['', ''],
        ['core/config.py', 'Thresholds, model names, environment loading'],
        ['core/llm.py', 'Gemini client with JSON repair'],
        ['core/agents/cv_agent.py', 'M1 — CV parsing'],
        ['core/agents/jd_agent.py', 'M2 — job description parsing'],
        ['core/agents/question_agent.py', 'M4 — question generation and graph-priority ordering'],
        ['core/graph/skill_graph.py', 'M3 — ESCO taxonomy, matching, gap analysis, graph payload'],
        ['core/graph/state.py', 'Per-skill status tracking across the interview'],
        ['core/graph/traversal.py', 'Decides which skills needed further probing'],
        ['core/livekit/run_agent.py', 'M5 — the voice interviewer agent'],
        ['core/livekit/launcher.py', 'Downloads and manages the LiveKit server process'],
        ['core/evaluator/evaluator.py', 'M6 — LLM-as-Judge answer scoring'],
        ['core/evaluator/integrity.py', 'M9 — Isolation Forest integrity detection'],
        ['core/evaluator/fusion.py', 'M11 — engagement and final weighted fusion'],
        ['core/agents/interviewer_prompt.py', 'The interviewer instructions shared by both modes'],
        ['core/pipeline/text_interview.py', 'Text mode - the typed interview engine'],
        ['core/pipeline/session_eval.py', 'Post-interview pipeline tying M6, M9, M11 and M12 together'],
        ['core/report/generator.py', 'M12 — final report assembly and reliability statistics'],
        ['data/esco/', 'ESCO taxonomy CSV files (1,201 digital skills)'],
        ['', ''],
        ['frontend/src/App.jsx', 'Six-step navigation; the interview takes the full viewport'],
        ['frontend/src/lib/vision.js', 'M7 and M8 — MediaPipe attention and posture'],
        ['frontend/src/lib/voice.js', 'M10 — Web Audio prosody analysis'],
        ['frontend/src/components/LandmarkOverlay.jsx', 'Draws the face mesh, gaze points and pose skeleton on the video'],
        ['frontend/src/screens/UploadStep.jsx', 'Step 1 — CV upload and JD paste'],
        ['frontend/src/screens/GraphStep.jsx', 'Step 2 — skill knowledge graph'],
        ['frontend/src/screens/QuestionsStep.jsx', 'Step 3 — generated questions'],
        ['frontend/src/screens/SetupScreen.jsx', 'Step 4 — mode choice, device check, briefing, prewarm trigger'],
        ['frontend/src/screens/InterviewScreen.jsx', 'Step 5 — the voice interview call UI'],
        ['frontend/src/screens/TextInterviewScreen.jsx', 'Step 5 — the typed interview chat UI'],
        ['frontend/src/screens/DashboardScreen.jsx', 'Step 6 — the assessment report'],
        ['frontend/scripts/fetch-vision-assets.mjs', 'Stages MediaPipe WASM and models at build time'],
    ],
    widths=[6.5, 10.0],
)


# ═════════════════════════════════════════════════════════════════════════
h1('7. API Endpoints')
# ═════════════════════════════════════════════════════════════════════════

table(
    ['Method and path', 'Purpose'],
    [
        ['POST /api/parse-cv', 'M1 — upload a CV as PDF or text, returns structured data'],
        ['POST /api/parse-jd', 'M2 — parse a job description'],
        ['POST /api/build-graph', 'M3 — build the skill graph, returns gaps, topics and the graph payload'],
        ['POST /api/generate-questions', 'M4 — generate the question set from graph topics'],
        ['POST /api/prewarm', 'Boot the media server and agent process early, from the setup screen'],
        ['POST /api/launch-interview', 'M5 — ensure the LiveKit server is up'],
        ['POST /api/text-interview/start', 'Text mode — begin a typed interview, returns the greeting'],
        ['POST /api/text-interview/answer', 'Text mode — submit an answer, returns the next question'],
        ['POST /api/text-interview/end', 'Text mode — close the interview at the candidate request'],
        ['GET  /token', 'Issue a LiveKit token and spawn the agent process'],
        ['POST /api/stop-interview', 'Shut down the agent process and LiveKit server'],
        ['POST /save_transcript', 'Transcript written by the client'],
        ['GET  /api/transcript', 'Most recently saved transcript'],
        ['POST /api/evaluate-session', 'M6+M9+M11+M12 — score a whole interview, returns the report'],
        ['POST /api/evaluate', 'M6 — score a single answer'],
        ['POST /api/integrity', 'M9 — integrity assessment for raw telemetry'],
        ['POST /api/fusion-report', 'M11 — fusion for supplied scores'],
        ['GET  /api/session', 'Current session state'],
        ['GET  /api/health', 'Health check and key configuration status'],
    ],
    widths=[6.0, 10.5],
)

para(
    'The report screen calls /api/evaluate-session automatically when it loads. That request '
    'runs three Gemini calls per scored answer (one reference answer, two rubric orderings) '
    'with four answers evaluated concurrently, so a ten-answer interview takes roughly 30 to '
    '60 seconds to score.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('8. Configuration and Tuning')
# ═════════════════════════════════════════════════════════════════════════

para('Everything worth adjusting is in one of these places.')

table(
    ['Setting', 'File', 'Default', 'Effect'],
    [
        ['MAX_INTERVIEW_QUESTIONS', '.env / launch request', '15', 'Question budget before wrap-up'],
        ['MIN_INTERVIEW_QUESTIONS', '.env', '5', 'Wrap-up will not trigger below this'],
        ['INTERVIEW_TIME_BUDGET_MINS', '.env / launch request', '30', 'Time budget before wrap-up'],
        ['GEMINI_MODEL', '.env', 'gemini-2.5-flash', 'Model used for all reasoning'],
        ['SCORE_STRONG_THRESHOLD', 'core/config.py', '70', 'Answer verdict band'],
        ['SCORE_WEAK_THRESHOLD', 'core/config.py', '40', 'Answer verdict band'],
        ['CONSISTENCY_HIGH / MODERATE', 'core/evaluator/evaluator.py', '8 / 16', 'Judge reliability bands'],
        ['WEIGHTS', 'core/evaluator/fusion.py', '.50/.20/.15/.15', 'Final fusion weighting'],
        ['ENGAGEMENT_WEIGHTS', 'core/evaluator/fusion.py', '.45/.20/.35', 'Attention / posture / voice mix'],
        ['NORMAL / SUSPICIOUS_THRESHOLD', 'core/evaluator/integrity.py', '60 / 35', 'Integrity verdict bands'],
        ['MIN_FUZZY_LEN, FUZZY_CUTOFF', 'core/graph/skill_graph.py', '6, 0.88', 'Skill matching strictness'],
        ['DISPLAY_LEAD_SECONDS', 'core/livekit/run_agent.py', '0.45', 'Delay between showing and speaking'],
        ['DETECT_INTERVAL_MS', 'frontend/src/lib/vision.js', '200', 'Landmark detection rate (overlay smoothness)'],
        ['YAW / PITCH_TOLERANCE', 'frontend/src/lib/vision.js', '0.42 / 0.38', 'Attention sensitivity'],
    ],
    widths=[5.0, 4.6, 3.0, 3.9],
)

para(
    'If the interview should be shorter for a demonstration, set MAX_INTERVIEW_QUESTIONS to '
    'about 6 and MIN_INTERVIEW_QUESTIONS to 3 in .env. The wrap-up watchdog checks every five '
    'seconds and will not trigger before the minimum is reached.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('9. Status, Limitations and Next Steps')
# ═════════════════════════════════════════════════════════════════════════

h2('9.1 What is complete')
bullet('CV and job description parsing into structured data')
bullet('ESCO skill graph with conservative matching, gap analysis and a clustered visualisation')
bullet('Graph-prioritised question generation')
bullet('Live voice interview with display-before-speak, budget enforcement and clean automatic ending')
bullet('Text interview mode sharing the same interviewer, budgets, monitoring and report')
bullet('Answer evaluation for every substantive answer, with rubric breakdown and self-consistency checking')
bullet('MediaPipe attention and posture analysis, calibrated per candidate')
bullet('Vocal delivery analysis from prosodic features')
bullet('Behavioural integrity detection, calibrated and always explained')
bullet('Weighted fusion into a final recommendation with full arithmetic exposed')
bullet('Report screen showing every component and its inputs')

h2('9.2 Known limitations')
para(
    'These are stated plainly because a handover that hides them is not useful.'
)
bullet('The trained-classifier track was removed after measurement showed the comparison could '
       'not be meaningful: its labels were LLM-generated, making agreement circular, and the '
       'trained model scored a correct paraphrase of the reference answer at 39 out of 100. '
       'The research question was narrowed to establishing the reliability of the judge itself. '
       'See docs/track-b-rejection.md.',
       bold_prefix='Second evaluation track — ')
bullet('No human-rated gold standard was collected, as that would have required ethical approval '
       'and participant recruitment. Agreement is therefore measured against answer-quality '
       'bands and between repeated judge runs, not against human raters.',
       bold_prefix='Human agreement study — ')
bullet('The judge scores against a reference answer that the same model family generated. This is '
       'reasonable for relative comparison but is not an independent ground truth.',
       bold_prefix='Reference answers — ')
bullet('The Isolation Forest baseline is synthetic, chosen to match the ranges the system measures. '
       'It has not been fitted to real pilot sessions, which the proposal intended.',
       bold_prefix='Integrity baseline — ')
bullet('Facial emotion still loads face-api.js and its models from a public CDN. If that is '
       'unreachable, emotion data is simply absent; the rest of the report is unaffected.',
       bold_prefix='Emotion detection — ')
bullet('Session state is held in memory in a single server process. It supports one interview at '
       'a time, which is appropriate for a dissertation demonstration but not for concurrent use.',
       bold_prefix='Single session — ')
bullet('The four rubric criteria are equally weighted at 25 points each, following the proposal. '
       'This means a fully accurate but partially complete answer can still score around 80.',
       bold_prefix='Equal rubric weights — ')

h2('9.3 Suggested next steps')
bullet('Run a small human-rating study over recorded answers and report quadratic-weighted '
       'Cohen\'s Kappa and Spearman correlation against the judge.')
bullet('Paraphrase a set of answers and measure score variance within each paraphrase group.')
bullet('Collect pilot sessions from volunteers and refit the integrity baseline on real data.')
bullet('Run the positional-bias ablation: score a set of answers under each rubric ordering '
       'alone and under the average, and test whether averaging measurably reduces the shift.')
bullet('Consider weighting technical accuracy and completeness above clarity and relevance, '
       'if partial answers should score lower.')

h2('9.4 Ethics')
para(
    'No real hiring decision is made with this system and no person\'s employment is affected '
    'by it. All evaluation data is synthetic or supplied by the person testing it. Video and '
    'audio are analysed in the browser and never recorded or transmitted; only derived numeric '
    'features leave the machine. Any pilot data collection involving human participants '
    'requires ethical approval, informed consent, anonymisation, and the right to withdraw.'
)


# ═════════════════════════════════════════════════════════════════════════
h1('10. Troubleshooting')
# ═════════════════════════════════════════════════════════════════════════

table(
    ['Symptom', 'Cause and fix'],
    [
        ['"python not found"', 'Reinstall Python with "Add to PATH" ticked'],
        ['"node not found"', 'Install Node.js LTS from nodejs.org'],
        ['Agent does not speak at all',
         'Every voice provider failed. agent_debug.log names each one and why. Most often both '
         'the ElevenLabs and Deepgram keys are missing, invalid, or out of credit'],
        ['Agent speaks a few words then goes silent',
         'Almost always an exhausted ElevenLabs quota — see section 5.7. The agent now detects '
         'this at startup and falls back to Deepgram automatically'],
        ['Header shows "text only"',
         'No voice provider produced audio. Questions are still displayed and the interview '
         'still works. Top up ElevenLabs or check the Deepgram key'],
        ['No attention or posture on the report',
         'The MediaPipe assets did not download. Run: cd frontend && npm run vision-assets'],
        ['No emotion data', 'face-api.js CDN was unreachable; everything else still works'],
        ['Camera or microphone blocked', 'Allow access via the lock icon in the browser address bar'],
        ['Port 8000 already in use', 'Close the previous server, or restart the machine'],
        ['Report says "Evaluation failed"',
         'Usually a Gemini rate limit or network error. Press Retry on the report screen'],
        ['Interview does not end', 'The watchdog force-closes 35 seconds after wrap-up is requested'],
    ],
    widths=[5.5, 11.0],
)

doc.add_paragraph()
para(
    'The agent writes a full log to agent_debug.log in the project root on every run. That '
    'file is the first place to look for anything related to the voice interview.',
    italic=True,
)


# ═════════════════════════════════════════════════════════════════════════
doc.save('PROJECT_DOCS.docx')
print('PROJECT_DOCS.docx written')
