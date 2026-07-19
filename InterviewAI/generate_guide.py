from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('InterviewAI — Complete Project Guide', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('A simple, clear explanation of the entire project').italic = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('CMP7200 Individual Masters Project').bold = True
doc.add_paragraph('Birmingham City University').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. WHAT IS THIS PROJECT?
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. What Is This Project?', level=1)
doc.add_paragraph(
    'InterviewAI is an AI-powered system that conducts technical interviews automatically. '
    'A candidate uploads their Resume and a Job Description, and the system:'
)
for s in [
    'Reads the Resume to find the candidate\'s skills and experience.',
    'Reads the Job Description to find what skills the job requires.',
    'Compares both using a standard skill database (ESCO taxonomy) to find gaps.',
    'Generates interview questions targeting the weak or missing skills.',
    'Asks questions one by one — either by typing (Text Mode) or by voice (LiveKit Mode).',
    'Scores each answer using AI (Google Gemini).',
    'Produces a final report: overall score, strengths, gaps, and hire decision.',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('You do NOT need any coding knowledge to use this. ').bold = True
p.add_run('Everything runs through a web browser.')

# ═══════════════════════════════════════════════════════════════════════════════
# 2. WHAT INPUTS DOES IT NEED?
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. What Inputs Does It Need?', level=1)

doc.add_paragraph('The system needs 3 things to work:')

doc.add_heading('1. Resume / CV (PDF or text)', level=2)
doc.add_paragraph('A file containing the candidate\'s work history, skills, and qualifications. '
                  'The system reads this to understand what the candidate already knows.')

doc.add_heading('2. Job Description (text)', level=2)
doc.add_paragraph('A description of the job role — what skills are required, what experience level '
                  '(junior/mid/senior), and what the role involves. The system reads this to know '
                  'what skills to test.')

doc.add_heading('3. API Keys (in .env file)', level=2)
doc.add_paragraph('Three API keys are required. These are stored in a file called .env:')
p = doc.add_paragraph()
p.add_run('GEMINI_API_KEY ').bold = True
p.add_run('— For all AI operations: reading CV/JD, generating questions, evaluating answers.')
p = doc.add_paragraph()
p.add_run('DEEPGRAM_API_KEY ').bold = True
p.add_run('— For voice mode only: converts speech to text (what the candidate says).')
p = doc.add_paragraph()
p.add_run('ELEVENLABS_API_KEY ').bold = True
p.add_run('— For voice mode only: converts text to speech (the AI asking questions out loud).')

doc.add_paragraph('')
doc.add_paragraph('For LiveKit Mode, the system also uses its own built-in LiveKit server for '
                   'managing the voice/video connection. This runs automatically — no extra setup needed.')

# ═══════════════════════════════════════════════════════════════════════════════
# 3. WHAT OUTPUTS DOES IT PRODUCE?
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. What Outputs Does It Produce?', level=1)
doc.add_paragraph('After the interview, the system produces a Final Report containing:')
for s in [
    'Overall Score — an average of all answers (out of 100).',
    'Verdict — Strong Hire (score ≥70), Weak Hire (40-69), or No Hire (<40).',
    'Per-Skill Breakdown — each skill tested, its score, and feedback.',
    'Strengths — skills where the candidate scored ≥70.',
    'Gaps — skills where the candidate scored <40.',
    'Development Areas — skills that need improvement (40-69).',
    'Full Answer Log — every question asked, the candidate\'s answer, the score, and the AI\'s feedback.',
]:
    doc.add_paragraph(s, style='List Bullet')

# ═══════════════════════════════════════════════════════════════════════════════
# 4. HOW IT WORKS — STEP BY STEP
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. How It Works — Step by Step', level=1)
doc.add_paragraph('The system runs in 6 stages, one after another:')

stages = [
    ('Stage 1: Parse CV & JD',
     'The candidate uploads their Resume and pastes the Job Description. The system sends both '
     'to Gemini AI, which extracts: the candidate\'s name and skills (from Resume), and the job '
     'title, level, and required skills (from JD). All of this is returned as structured data.'),

    ('Stage 2: Build Skill Graph',
     'The system loads the ESCO taxonomy — a European-standard database of 1,201 IT skills — into '
     'a graph (like a map where skills are connected to related skills). It adds modern technologies '
     'like Python, React, and Docker that ESCO doesn\'t cover. The candidate\'s skills and the job\'s '
     'required skills are matched against this graph.'),

    ('Stage 3: Gap Analysis',
     'The system compares what the candidate knows vs what the job needs. Skills that are missing '
     'or weak become "topics" for the interview. Each topic gets a priority level and a reason.'),

    ('Stage 4: Generate Questions',
     'Using Gemini AI, the system creates a full interview question set: opening questions (warm-up), '
     'technical questions (one per skill to test), behavioural questions (STAR format), and closing '
     'questions. Each question has a difficulty level and a follow-up probe.'),

    ('Stage 5: Conduct Interview',
     'The interview runs in one of two modes:\n'
     '• Text Mode: Questions appear on screen. The candidate types answers and clicks Submit. '
     'Each answer is immediately scored by Gemini AI (two evaluations averaged for accuracy). '
     'The system adapts — it picks the weakest skill first and asks follow-ups if needed.\n'
     '• LiveKit Mode: A voice agent speaks the questions using ElevenLabs (AI voice). The candidate '
     'speaks their answers. Deepgram converts speech to text. Everything happens in real time through '
     'a web page. At the end, a transcript is saved.'),

    ('Stage 6: Generate Report',
     'After the interview ends, the system builds the final report: overall score, verdict, per-skill '
     'breakdown, strengths, gaps, development areas, and full answer log.'),
]

for title, desc in stages:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. THE TWO INTERVIEW MODES (EXPLAINED SIMPLY)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. The Two Interview Modes', level=1)

doc.add_heading('Text Mode', level=2)
doc.add_paragraph(
    'This is the simpler mode. The candidate sees questions on a web page, types their answer in a '
    'text box, and clicks Submit. The AI scores it immediately and moves to the next question. '
    'No microphone or camera needed.'
)
doc.add_paragraph('Flow:')
for s in [
    'System picks the lowest-scored skill first.',
    'Generates a fresh question for that skill using Gemini.',
    'Displays the question on screen.',
    'Candidate types answer and clicks Submit.',
    'AI evaluates the answer (two Gemini calls, scores averaged).',
    'Result shown immediately. If weak, a follow-up may be asked.',
    'Repeats until all skills tested or max questions reached.',
    'Final report generated automatically.',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('LiveKit Mode (Voice Interview)', level=2)
doc.add_paragraph(
    'In this mode, the candidate speaks their answers through a microphone and the AI speaks the '
    'questions out loud. It works like a real voice call.'
)
doc.add_paragraph('What is LiveKit?')
doc.add_paragraph(
    'LiveKit is an open-source platform for real-time voice/video communication (similar to Zoom). '
    'It runs as a server on your machine. The system starts it automatically when you use voice mode.'
)
doc.add_paragraph('How the Voice Interview Works:')
for s in [
    '1. The system starts a small web server (port 18765) and the LiveKit server (port 7880).',
    '2. The candidate\'s browser opens a web page (client.html) that shows their webcam and a chat log.',
    '3. An AI "agent" (run_agent.py) joins the same room as the candidate.',
    '4. The agent speaks pre-generated questions one by one using ElevenLabs TTS (text-to-speech).',
    '   ElevenLabs converts text into natural-sounding human speech.',
    '5. The candidate hears the question and speaks their answer into the microphone.',
    '6. Deepgram STT (speech-to-text) converts their speech into text in real time.',
    '7. The text is displayed on the screen as "live transcription."',
    '8. After all 5 questions, the agent says goodbye and saves a transcript file.',
    '9. The candidate can end the interview anytime by clicking End.',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Key point: ').bold = True
p.add_run(
    'The LiveKit agent does NOT use Gemini during the conversation. All 5 questions are pre-loaded '
    'from the question generation stage. Gemini is only used before the interview starts (to create '
    'the questions). This makes the voice interview faster and more reliable.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MODULES — WHAT EACH PART DOES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Modules — What Each Part Does', level=1)
doc.add_paragraph('The project has 7 modules (plus configuration and AI client). Each module has a specific job:')

modules = [
    ('Module 1: agents/',
     'Parsing & Question Generation',
     'This module handles understanding the inputs and creating questions.\n'
     '• cv_agent.py — Takes a Resume PDF or text, extracts name + skills + experience using Gemini.\n'
     '• jd_agent.py — Takes a Job Description, extracts job title + level + required skills using Gemini.\n'
     '• question_agent.py — Takes the skill analysis topics and generates full interview questions using Gemini.'),

    ('Module 2: evaluator/',
     'Answer Scoring (LLM-as-Judge)',
     'This module scores the candidate\'s answers.\n'
     '• evaluator.py — For each answer, it first generates an "ideal answer" using Gemini, then compares '
     'the candidate\'s answer against it. Scores on 4 criteria (technical accuracy, completeness, clarity, '
     'relevance — each 0-25, total 0-100). Runs two evaluations with different criterion order to reduce '
     'AI bias and averages them.'),

    ('Module 3: graph/',
     'Skill Knowledge Graph & State Tracking',
     'This module manages the skill database and tracks interview progress.\n'
     '• skill_graph.py — Builds a map of 1,201 IT skills from ESCO taxonomy CSV files using NetworkX. '
     'Adds modern tech not in ESCO. Matches skills using fuzzy logic.\n'
     '• state.py — Tracks each skill\'s status during the interview: pending, verified strong, verified weak, '
     'confirmed gap. Records scores and decides when a skill is fully tested.\n'
     '• traversal.py — Decides which skill to ask next. Picks the lowest-scored skill first (adaptive strategy). '
     'Also decides if a follow-up question is needed for weak answers.\n'
     '• visualize.py — Creates charts showing candidate skills, job requirements, gaps, and full skill map.'),

    ('Module 4: pipeline/',
     'Interview Orchestrator',
     'This is the "brain" that runs the text interview.\n'
     '• interview_loop.py — The InterviewLoop class runs the question → answer → evaluate cycle. '
     'Picks the next skill, generates a question, collects the answer, scores it, updates progress, '
     'and repeats until done.'),

    ('Module 5: report/',
     'Final Report Generator',
     'This module creates the final output after the interview.\n'
     '• generator.py — Takes all answers and scores, computes overall average, assigns verdict '
     '(Strong Hire ≥70, Weak Hire 40-69, No Hire <40), lists strengths and gaps, and builds '
     'the complete report with answer log.'),

    ('Module 6: livekit/',
     'Voice Interview System',
     'This module handles everything related to voice interviews.\n'
     '• run_agent.py — The AI voice agent. Speaks questions (ElevenLabs), listens to answers (Deepgram), '
     'publishes live transcript to the web page.\n'
     '• whisper_server.py — A tiny web server that serves the voice interview web page and generates '
     'LiveKit access tokens.\n'
     '• launcher.py — Starts and manages the LiveKit server.\n'
     '• client.html — The web page the candidate sees: webcam, transcription, chat log.\n'
     '• livekit.yaml — Configuration file for the LiveKit server.'),

    ('Module 7: config.py + llm.py',
     'Settings & AI Communication',
     'These two files support everything else.\n'
     '• config.py — Loads .env file, defines all constants (scoring thresholds, question limits, '
     'API keys, model names). One central place to change settings.\n'
     '• llm.py — Talks to Google Gemini API. Has two functions: call_llm() gets text responses, '
     'call_llm_json() gets JSON responses (auto-fixes broken JSON).'),
]

for emoji, title, desc in modules:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. DIRECTORY STRUCTURE — EVERY FILE EXPLAINED
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Directory Structure — Every File Inside core/ Explained', level=1)
doc.add_paragraph(
    'Below is every file and folder inside the core/ directory. Each one is explained in simple terms.'
)

files = [
    ('core/config.py',
     'Settings file. Loads the .env file and defines all the numbers and rules the project uses: '
     'Gemini API key and model, scoring thresholds (70 = strong pass, 40 = weak pass), max questions '
     'per session (3), max follow-ups (1), and status labels for the skill graph. Every other file '
     'imports from here so settings are in one place.'),

    ('core/llm.py',
     'AI communication module. Contains the code that calls Google Gemini API. Two main functions: '
     'call_llm() sends text and gets text back. call_llm_json() sends text and expects JSON back — '
     'it automatically removes markdown formatting and fixes broken JSON using json_repair library. '
     'Used by question generator, evaluator, and all agents.'),

    ('core/agents/ (folder — 3 files)',
     'This folder handles input parsing and question creation.\n\n'
     'core/agents/cv_agent.py — Reads the Resume. If PDF, extracts text first. Then sends to Gemini '
     'to extract: candidate name, skills list, experience summary.\n\n'
     'core/agents/jd_agent.py — Reads the Job Description. Sends to Gemini to extract: job title, '
     'role level (junior/mid/senior), required skills list.\n\n'
     'core/agents/question_agent.py — Creates interview questions. Three functions: '
     'generate_interview_questions() creates the full structured question set from skill topics. '
     'generate_position_question() creates one question on the fly during the live interview. '
     'build_interview_flow() flattens the structured set into a simple ordered list.'),

    ('core/evaluator/ (folder — 1 file)',
     'This folder handles scoring the candidate\'s answers.\n\n'
     'core/evaluator/evaluator.py — The scoring engine. For each answer: generates an "ideal" reference '
     'answer using Gemini, then compares the real answer against it. Scores on 4 criteria each worth '
     '25 points: technical accuracy, completeness, clarity, relevance. To be fair, it runs the '
     'evaluation twice with criteria in different order and averages the scores. Answers shorter '
     'than 10 words automatically score 0. Returns score (0-100), verdict, and feedback.'),

    ('core/graph/ (folder — 4 files)',
     'This folder manages the skill database and tracks interview progress.\n\n'
     'core/graph/skill_graph.py — Builds a skill map from ESCO taxonomy CSVs (1,201 IT skills) '
     'using NetworkX. Adds modern tech manually (Python, React, Docker, etc.). Can fuzzy-match '
     'skill names. The gap_analysis() function finds what skills the candidate is missing.\n\n'
     'core/graph/state.py — Tracks progress during the interview. Each skill has a node that stores: '
     'status (pending/verified/gap), scores, questions asked, best score. record_answer() updates '
     'status based on score. InterviewState tracks all skills together.\n\n'
     'core/graph/traversal.py — Decides which skill to ask next. Picks the lowest-scored skill first, '
     'prioritizing high-priority skills, avoiding repeats. decide_follow_up() checks if a weak answer '
     'needs a follow-up question.\n\n'
     'core/graph/visualize.py — Draws charts showing candidate skills, job requirements, skill gaps, '
     'and the full skill map. These appear in the Streamlit UI.'),

    ('core/pipeline/ (folder — 1 file)',
     'This folder orchestrates the text interview.\n\n'
     'core/pipeline/interview_loop.py — The InterviewLoop class runs the interview cycle: '
     'get_next_question() picks the next skill and generates a fresh question. submit_answer() '
     'scores the answer and updates state. Keeps going until all skills verified or max questions '
     'reached. get_summary() collects everything for the final report.'),

    ('core/report/ (folder — 1 file)',
     'This folder creates the final report after the interview.\n\n'
     'core/report/generator.py — Takes the InterviewLoop data, computes overall score, assigns '
     'verdict (Strong Hire ≥70, Weak Hire 40-69, No Hire <40), lists strengths and gaps, and '
     'builds the complete report dictionary with answer log.'),

    ('core/livekit/ (folder — 4 files + client.html + livekit.yaml)',
     'This folder handles the voice interview mode.\n\n'
     'core/livekit/run_agent.py — The AI voice agent. Connects to LiveKit room, speaks pre-loaded '
     'questions using ElevenLabs TTS, listens to answers via Deepgram STT, publishes live '
     'transcript to the web page. Saves transcript JSON after all questions.\n\n'
     'core/livekit/whisper_server.py — Tiny web server (port 18765). Serves the voice interview '
     'web page (client.html) at /livekit. Generates LiveKit tokens at /token. Saves transcripts '
     'at /save_transcript. Runs in a background thread.\n\n'
     'core/livekit/launcher.py — Manages the LiveKit server. start_livekit_server() launches it '
     'if not running. launch() starts everything and returns the client URL. cleanup() shuts '
     'everything down on exit.\n\n'
     'core/livekit/client.html — The web page the candidate sees during voice interview. Shows '
     'webcam feed, real-time transcription, chat conversation log, and current question. Auto-'
     'connects to LiveKit on load.\n\n'
     'core/livekit/livekit.yaml — Configuration for the LiveKit server: ports, API keys, logging.'),
]

for filepath, desc in files:
    p = doc.add_paragraph()
    run = p.add_run(filepath)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(desc)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Summary', level=1)
doc.add_paragraph(
    'InterviewAI takes a Resume + Job Description → parses them → builds a skill graph → '
    'finds gaps → generates interview questions → conducts the interview (text or voice) → '
    'scores each answer using AI → produces a final report with verdict and recommendations.'
)
doc.add_paragraph('')
doc.add_paragraph('Three API keys power everything:')
p = doc.add_paragraph()
p.add_run('Gemini').bold = True
p.add_run(' — reads CV/JD, generates questions, evaluates answers')
p = doc.add_paragraph()
p.add_run('Deepgram').bold = True
p.add_run(' — converts speech to text (voice mode only)')
p = doc.add_paragraph()
p.add_run('ElevenLabs').bold = True
p.add_run(' — converts text to speech (voice mode only)')
doc.add_paragraph('')
doc.add_paragraph(
    'All the logic lives inside core/. Seven modules handle: settings (config.py), AI communication '
    '(llm.py), input parsing (agents/), answer scoring (evaluator/), skill knowledge (graph/), '
    'interview orchestration (pipeline/), report generation (report/), and voice interviews (livekit/).'
)

# ── Save ──
out = Path(__file__).resolve().parent / 'PROJECT_GUIDE.docx'
doc.save(out)
print(f'Saved to {out}')
