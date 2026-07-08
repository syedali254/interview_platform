"""
CMP7200 Individual Masters Project - Final Proposal Generator (v3)
Generates: 5 diagrams (PNG) + 1 Word document (.docx)

CMP7200 formatting requirements:
- Font: Arial, size 11
- Line spacing: 1.5
- Justified text
- Anonymous marking (student number only)
- BCU Harvard referencing

Content for 80-90% band:
- Conceptual model from literature
- Alternative methods considered
- Limitations exhaustively identified
- Supporting appendices
"""

import os
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Circle
from matplotlib.lines import Line2D
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(os.path.abspath(__file__))

# ─── Colours ──────────────────────────────────────────────────────────────────
NAV = '#1b2838'
BLU = '#2c5282'
STL = '#4a7fb5'
TEA = '#1a7a6d'
GRN = '#38805b'
AMB = '#c05621'
SLT = '#4a5568'
GRY = '#a0aec0'
LGT = '#f7fafc'
WHT = '#ffffff'
DRK = '#1a202c'


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 1: SYSTEM ARCHITECTURE — wide, generous spacing
# ══════════════════════════════════════════════════════════════════════════════

def build_architecture():
    fig, ax = plt.subplots(figsize=(18, 13))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 13)
    ax.axis('off')
    fig.patch.set_facecolor(WHT)

    def box(x, y, w, h, c, txt, sub=''):
        sh = FancyBboxPatch((x+.06, y-.06), w, h,
            boxstyle="round,pad=0.12", fc='#00000012', ec='none', zorder=1)
        ax.add_patch(sh)
        b = FancyBboxPatch((x, y), w, h,
            boxstyle="round,pad=0.12", fc=c, ec='#00000022',
            lw=1.2, alpha=.92, zorder=2)
        ax.add_patch(b)
        if sub:
            ax.text(x+w/2, y+h/2+.22, txt, ha='center', va='center',
                fontsize=12, color='white', fontweight='bold', fontfamily='Verdana', zorder=3)
            ax.text(x+w/2, y+h/2-.22, sub, ha='center', va='center',
                fontsize=9, color='#ffffffcc', fontfamily='Verdana', zorder=3)
        else:
            ax.text(x+w/2, y+h/2, txt, ha='center', va='center',
                fontsize=12, color='white', fontweight='bold', fontfamily='Verdana',
                zorder=3, linespacing=1.3)

    def phase(y, h, label, c):
        bg = FancyBboxPatch((1.5, y), 16, h,
            boxstyle="round,pad=0.15", fc=c, ec='#e2e8f0',
            lw=.8, alpha=.08, zorder=0)
        ax.add_patch(bg)
        bar = FancyBboxPatch((.3, y+.15), 1.0, h-.3,
            boxstyle="round,pad=0.06", fc=c, ec='none', alpha=.85, zorder=2)
        ax.add_patch(bar)
        ax.text(.8, y+h/2, label, ha='center', va='center',
            fontsize=10.5, color='white', fontweight='bold',
            fontfamily='Verdana', rotation=90, zorder=3)

    def arrow(x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=SLT, lw=2))

    # Phase bands
    phase(9.5, 2.8, 'PHASE 1', BLU)
    phase(6.2, 2.8, 'PHASE 2', STL)
    phase(3.3, 2.4, 'PHASE 3', TEA)
    phase(.5, 2.3, 'PHASE 4', AMB)

    # Phase labels (right side)
    ax.text(17.3, 10.9, 'Pre-Interview', ha='right', fontsize=10, color=BLU,
        fontweight='bold', fontfamily='Verdana', style='italic')
    ax.text(17.3, 7.6, 'Live Interview', ha='right', fontsize=10, color=STL,
        fontweight='bold', fontfamily='Verdana', style='italic')
    ax.text(17.3, 4.5, 'Evaluation', ha='right', fontsize=10, color=TEA,
        fontweight='bold', fontfamily='Verdana', style='italic')
    ax.text(17.3, 1.65, 'Reporting', ha='right', fontsize=10, color=AMB,
        fontweight='bold', fontfamily='Verdana', style='italic')

    # ── Phase 1 modules ──
    box(2.0, 10.0, 3.2, 1.8, BLU, 'M1: CV Parser', 'LLM Agent')
    box(5.8, 10.0, 3.2, 1.8, BLU, 'M2: JD Analyser', 'LLM Agent')
    box(9.6, 10.0, 3.2, 1.8, BLU, 'M3: Skill Graph', 'NetworkX + ESCO')
    box(13.4, 10.0, 3.2, 1.8, BLU, 'M4: Question Gen', 'LLM + Graph')

    # ── Phase 2 modules ──
    box(1.8, 6.7, 2.8, 1.8, STL, 'M5: Voice', 'Whisper + TTS')
    box(5.0, 6.7, 2.8, 1.8, STL, 'M7: Vision', 'MediaPipe')
    box(8.2, 6.7, 2.8, 1.8, STL, 'M8: Posture', 'MediaPipe Pose')
    box(11.4, 6.7, 2.8, 1.8, STL, 'M9: Integrity', 'Isolation Forest')
    box(14.6, 6.7, 2.8, 1.8, STL, 'M10: Emotion', 'wav2vec2')

    # ── Phase 3: Core research ──
    hl = FancyBboxPatch((2.0, 3.5), 13.6, 2.0,
        boxstyle="round,pad=0.15", fc='none', ec=TEA,
        lw=2.5, ls='--', zorder=1)
    ax.add_patch(hl)
    ax.text(8.8, 5.65, '[ CORE RESEARCH CONTRIBUTION ]', ha='center',
        fontsize=10, color=TEA, fontweight='bold', fontfamily='Verdana', style='italic')

    box(2.5, 3.7, 5.5, 1.6, TEA, 'M6-A: LLM-as-Judge', 'GPT-4 / Llama 3 Rubric Scoring')
    box(9.5, 3.7, 5.5, 1.6, '#0e6655', 'M6-B: ML Classifier', 'S-BERT + XGBoost + SHAP')

    # ── Phase 4 ──
    box(3.0, .8, 5.0, 1.6, AMB, 'M11: Recommender', 'Weighted Score Fusion')
    box(9.5, .8, 5.0, 1.6, AMB, 'M12: Report Gen', 'Template + Final Scores')

    # ── Arrows (simplified — one per phase transition) ──
    arrow(3.6, 10.0, 3.2, 8.5)
    arrow(7.4, 10.0, 6.4, 8.5)
    arrow(11.2, 10.0, 9.6, 8.5)
    arrow(15.0, 10.0, 16.0, 8.5)

    for x in [3.2, 6.4, 9.6, 12.8, 16.0]:
        arrow(x, 6.7, 8.8, 5.7)

    arrow(5.25, 3.7, 5.5, 2.4)
    arrow(12.25, 3.7, 12.0, 2.4)

    # Title
    ax.text(9, 12.7, 'System Architecture — Multi-Agent AI Interview Platform',
        ha='center', fontsize=16, color=DRK, fontweight='bold', fontfamily='Verdana')

    # Legend
    leg = [
        mpatches.Patch(color=BLU, alpha=.9, label='Phase 1: Pre-Interview'),
        mpatches.Patch(color=STL, alpha=.9, label='Phase 2: Live Interview'),
        mpatches.Patch(color=TEA, alpha=.9, label='Phase 3: Evaluation (Core)'),
        mpatches.Patch(color=AMB, alpha=.9, label='Phase 4: Reporting'),
    ]
    ax.legend(handles=leg, loc='lower right', fontsize=11, framealpha=.95,
        edgecolor='#e2e8f0', fancybox=True)

    p = os.path.join(OUT, 'final_architecture.png')
    plt.savefig(p, dpi=180, bbox_inches='tight', facecolor=WHT, pad_inches=0.3)
    plt.close()
    print(f'[1/5] Architecture: {p}')
    return p


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 2: USE CASE — extra wide, big text
# ══════════════════════════════════════════════════════════════════════════════

def build_usecase():
    fig, ax = plt.subplots(figsize=(16, 12))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 12)
    ax.axis('off')
    fig.patch.set_facecolor(WHT)

    def actor(x, y, label):
        head = Circle((x, y+.8), .25, fc=LGT, ec=DRK, lw=2.2, zorder=5)
        ax.add_patch(head)
        ax.plot([x,x], [y+.55, y-.08], color=DRK, lw=2.2, zorder=5)
        ax.plot([x-.3, x+.3], [y+.3, y+.3], color=DRK, lw=2.2, zorder=5)
        ax.plot([x, x-.22], [y-.08, y-.5], color=DRK, lw=2.2, zorder=5)
        ax.plot([x, x+.22], [y-.08, y-.5], color=DRK, lw=2.2, zorder=5)
        ax.text(x, y-.85, label, ha='center', fontsize=13, color=DRK,
            fontweight='bold', fontfamily='Verdana')

    def uc(x, y, txt, c=STL):
        sh = mpatches.Ellipse((x+.04, y-.04), 4.2, 1.15,
            fc='#00000008', ec='none', zorder=1)
        ax.add_patch(sh)
        e = mpatches.Ellipse((x, y), 4.2, 1.15,
            fc=c, ec=DRK, alpha=.12, lw=1.5, zorder=2)
        ax.add_patch(e)
        b = mpatches.Ellipse((x, y), 4.2, 1.15,
            fc='none', ec=c, lw=2.5, zorder=3)
        ax.add_patch(b)
        ax.text(x, y, txt, ha='center', va='center', fontsize=11,
            color=DRK, fontfamily='Verdana', fontweight='semibold', zorder=4)

    def line(x1, y1, x2, y2, ls='-'):
        ax.plot([x1,x2], [y1,y2], ls=ls, color=GRY, lw=1.5, zorder=1)

    # Title
    ax.text(8, 11.7, 'Use Case Diagram — AI Interview Platform',
        ha='center', fontsize=15, color=DRK, fontweight='bold', fontfamily='Verdana')

    # System boundary
    bd = FancyBboxPatch((2.8, .5), 10.4, 10.5,
        boxstyle="round,pad=0.25", fc='#f7fafc', ec=SLT,
        lw=2, ls='--', alpha=.5, zorder=0)
    ax.add_patch(bd)
    ax.text(8, 10.7, 'Interview Platform System', ha='center',
        fontsize=11, color=SLT, style='italic', fontfamily='Verdana')

    # Actors
    actor(1.0, 7.0, 'Candidate')
    actor(15.0, 7.0, 'Recruiter')

    # Use cases — generous vertical spacing
    cases = [
        (8, 9.8, 'Upload CV and Job Description', BLU),
        (8, 8.4, 'Build Candidate Skill Graph', BLU),
        (8, 7.0, 'Conduct Voice Interview', STL),
        (8, 5.6, 'Evaluate Answers (LLM + ML)', TEA),
        (8, 4.2, 'Monitor Behaviour and Integrity', TEA),
        (8, 2.8, 'Detect Emotion and Posture', STL),
        (8, 1.4, 'Generate Report and Recommendation', AMB),
    ]
    for cx, cy, ct, cc in cases:
        uc(cx, cy, ct, cc)

    # Candidate lines (top 4)
    for cx, cy, _, _ in cases[:4]:
        line(1.5, 7.0, cx-2.1, cy)

    # Recruiter lines (bottom 3 + evaluation)
    line(14.5, 7.0, cases[3][0]+2.1, cases[3][1])
    for cx, cy, _, _ in cases[4:]:
        line(14.5, 7.0, cx+2.1, cy)

    # Include
    line(cases[0][0], cases[0][1]-.58, cases[1][0], cases[1][1]+.58, ls='--')
    ax.text(6.5, 9.1, '\u00abinclude\u00bb', fontsize=9, color=BLU,
        style='italic', fontfamily='Verdana')

    p = os.path.join(OUT, 'final_usecase.png')
    plt.savefig(p, dpi=180, bbox_inches='tight', facecolor=WHT, pad_inches=0.3)
    plt.close()
    print(f'[2/5] Use Case: {p}')
    return p


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 3: MODULE 6 PIPELINE — wide dual-track
# ══════════════════════════════════════════════════════════════════════════════

def build_pipeline():
    fig, ax = plt.subplots(figsize=(16, 8))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 8)
    ax.axis('off')
    fig.patch.set_facecolor(WHT)

    def box(x, y, w, h, c, txt, sub=''):
        sh = FancyBboxPatch((x+.05, y-.05), w, h,
            boxstyle="round,pad=0.1", fc='#00000010', ec='none', zorder=1)
        ax.add_patch(sh)
        b = FancyBboxPatch((x, y), w, h,
            boxstyle="round,pad=0.1", fc=c, ec='#00000020',
            lw=1, alpha=.92, zorder=2)
        ax.add_patch(b)
        if sub:
            ax.text(x+w/2, y+h/2+.22, txt, ha='center', va='center',
                fontsize=12, color='white', fontweight='bold', fontfamily='Verdana', zorder=3)
            ax.text(x+w/2, y+h/2-.22, sub, ha='center', va='center',
                fontsize=9, color='#ffffffbb', fontfamily='Verdana', zorder=3)
        else:
            ax.text(x+w/2, y+h/2, txt, ha='center', va='center',
                fontsize=12, color='white', fontweight='bold', fontfamily='Verdana',
                zorder=3, linespacing=1.2)

    def arrow(x1, y1, x2, y2, c=SLT):
        ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
            arrowprops=dict(arrowstyle='->', color=c, lw=2.5))

    # Title
    ax.text(8, 7.6, 'Module 6 — Answer Evaluation Pipeline',
        ha='center', fontsize=14, color=DRK, fontweight='bold', fontfamily='Verdana')
    ax.text(8, 7.2, '(Core Research Contribution)', ha='center',
        fontsize=11, color=TEA, fontfamily='Verdana', style='italic')

    # Track backgrounds
    ta = FancyBboxPatch((.8, 5.0), 11.5, 2.0,
        boxstyle="round,pad=0.12", fc='#2c528208', ec='#2c528225',
        lw=1.2, ls='--', zorder=0)
    ax.add_patch(ta)
    tb = FancyBboxPatch((.8, 1.0), 11.5, 2.0,
        boxstyle="round,pad=0.12", fc='#1a7a6d08', ec='#1a7a6d25',
        lw=1.2, ls='--', zorder=0)
    ax.add_patch(tb)

    # Track labels
    ax.text(.6, 6.8, 'TRACK A', fontsize=10, color=BLU,
        fontweight='bold', fontfamily='Verdana', style='italic')
    ax.text(.6, 2.8, 'TRACK B', fontsize=10, color=TEA,
        fontweight='bold', fontfamily='Verdana', style='italic')

    # INPUT
    box(.3, 3.3, 2.5, 1.6, SLT, 'Candidate\nAnswer')

    # Track A
    box(3.5, 5.3, 3.3, 1.4, BLU, 'LLM Judge', 'GPT-4 / Llama 3')
    box(7.5, 5.3, 3.3, 1.4, STL, 'Rubric Scoring', '4 Criteria')
    box(11.5, 5.3, 1.8, 1.4, '#34568B', 'Score A')

    # Track B
    box(3.5, 1.3, 3.3, 1.4, TEA, 'Feature Extract', 'S-BERT + NLP')
    box(7.5, 1.3, 3.3, 1.4, GRN, 'XGBoost', 'Classifier')
    box(11.5, 1.3, 1.8, 1.4, '#0e6655', 'Score B')

    # FUSION
    box(13.0, 3.0, 2.6, 2.0, AMB, 'Compare\nand Fuse')

    # Arrows
    arrow(2.8, 4.5, 3.5, 6.0, BLU)
    arrow(2.8, 3.7, 3.5, 2.0, TEA)
    arrow(6.8, 6.0, 7.5, 6.0, BLU)
    arrow(10.8, 6.0, 11.5, 6.0, BLU)
    arrow(6.8, 2.0, 7.5, 2.0, TEA)
    arrow(10.8, 2.0, 11.5, 2.0, TEA)
    arrow(13.0, 5.8, 13.0, 5.0, BLU)
    arrow(13.0, 2.2, 13.0, 3.0, TEA)

    # Output arrow
    ax.annotate('FINAL\nSCORE', xy=(16, 4.0), fontsize=11, color=DRK,
        fontweight='bold', fontfamily='Verdana', ha='center', va='center')
    arrow(15.6, 4.0, 15.8, 4.0, AMB)

    # SHAP note
    ax.text(9.2, .6, '+ SHAP Explanations for every prediction',
        fontsize=10, color=TEA, fontweight='bold', fontfamily='Verdana', style='italic')

    p = os.path.join(OUT, 'final_pipeline.png')
    plt.savefig(p, dpi=180, bbox_inches='tight', facecolor=WHT, pad_inches=0.3)
    plt.close()
    print(f'[3/5] Pipeline: {p}')
    return p


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 4: DATA FLOW — step-by-step interview process
# ══════════════════════════════════════════════════════════════════════════════

def build_dataflow():
    fig, ax = plt.subplots(figsize=(16, 8))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 8)
    ax.axis('off')
    fig.patch.set_facecolor(WHT)

    def step(x, y, w, h, c, num, txt):
        sh = FancyBboxPatch((x+.04, y-.04), w, h,
            boxstyle="round,pad=0.08", fc='#00000010', ec='none', zorder=1)
        ax.add_patch(sh)
        b = FancyBboxPatch((x, y), w, h,
            boxstyle="round,pad=0.08", fc=c, ec='#00000018',
            lw=.8, alpha=.9, zorder=2)
        ax.add_patch(b)
        # Number badge
        badge = Circle((x+.45, y+h-.35), .25, fc='white', ec=c, lw=2, zorder=4)
        ax.add_patch(badge)
        ax.text(x+.45, y+h-.35, str(num), ha='center', va='center',
            fontsize=10, color=c, fontweight='bold', fontfamily='Verdana', zorder=5)
        ax.text(x+w/2, y+h/2-.15, txt, ha='center', va='center',
            fontsize=11, color='white', fontweight='bold', fontfamily='Verdana',
            zorder=3, linespacing=1.2)

    def arrow(x1, y1, x2, y2):
        ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
            arrowprops=dict(arrowstyle='->', color=SLT, lw=2))

    # Title
    ax.text(8, 7.6, 'Interview Process — Data Flow Diagram',
        ha='center', fontsize=14, color=DRK, fontweight='bold', fontfamily='Verdana')

    # Top row: Steps 1-4
    step(.5, 4.3, 3.2, 2.2, BLU, 1, 'CV + JD\nUpload')
    step(4.3, 4.3, 3.2, 2.2, BLU, 2, 'Skill Graph\nBuilt')
    step(8.1, 4.3, 3.2, 2.2, BLU, 3, 'Questions\nGenerated')
    step(11.9, 4.3, 3.2, 2.2, STL, 4, 'Voice\nInterview')

    # Bottom row: Steps 5-7
    step(.8, .8, 3.5, 2.2, TEA, 5, 'Answer\nEvaluation')
    step(5.0, .8, 3.5, 2.2, TEA, 6, 'Behaviour\nAnalysis')
    step(9.2, .8, 3.5, 2.2, AMB, 7, 'Final Report\nGenerated')

    # Output
    step(13.4, .8, 2.2, 2.2, '#6b3a2a', 8, 'Decision')

    # Arrows between top row
    arrow(3.7, 5.4, 4.3, 5.4)
    arrow(7.5, 5.4, 8.1, 5.4)
    arrow(11.3, 5.4, 11.9, 5.4)

    # Down from step 4 to step 5
    arrow(13.5, 4.3, 2.55, 3.0)

    # Arrows between bottom row
    arrow(4.3, 1.9, 5.0, 1.9)
    arrow(8.5, 1.9, 9.2, 1.9)
    arrow(12.7, 1.9, 13.4, 1.9)

    # Data labels
    ax.text(3.0, 6.7, 'Skills + Requirements', fontsize=9, color=GRY,
        fontfamily='Verdana', style='italic')
    ax.text(6.8, 6.7, 'Adaptive Questions', fontsize=9, color=GRY,
        fontfamily='Verdana', style='italic')
    ax.text(10.6, 6.7, 'Voice Responses', fontsize=9, color=GRY,
        fontfamily='Verdana', style='italic')
    ax.text(5.5, .4, 'Scores + Flags', fontsize=9, color=GRY,
        fontfamily='Verdana', style='italic')

    p = os.path.join(OUT, 'final_dataflow.png')
    plt.savefig(p, dpi=180, bbox_inches='tight', facecolor=WHT, pad_inches=0.3)
    plt.close()
    print(f'[4/5] Data Flow: {p}')
    return p


# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAM 5: GANTT CHART
# ══════════════════════════════════════════════════════════════════════════════

def build_gantt():
    fig, ax = plt.subplots(figsize=(16, 8.5))

    tasks = [
        ('Literature Review and Dataset Collection', 1, 2, BLU),
        ('Skill Graph and Parsing Agents (M1-M3)', 3, 1.5, BLU),
        ('Question Generator Module (M4)', 4, 1, BLU),
        ('Voice Interview Agent (M5)', 5, 1.5, STL),
        ('Multimodal Sensors (M7, M8, M10)', 5, 2, STL),
        ('LLM-as-Judge Baseline (M6 Track A)', 6, 2, TEA),
        ('Trained ML Classifier (M6 Track B)', 8, 2, TEA),
        ('Behavioural Integrity Model (M9)', 9, 1, TEA),
        ('System Integration and Testing', 10, 1.5, GRN),
        ('Comparison Experiments and Analysis', 10, 2, GRN),
        ('Write Chapters 1-3', 11, 1.5, AMB),
        ('Write Chapters 4-6', 12, 1.5, AMB),
        ('Final Draft and Proofreading', 13, 1, AMB),
        ('Viva Slides and Rehearsal', 14, 1, AMB),
    ]

    yp = list(range(len(tasks), 0, -1))
    for i, (nm, st, dur, c) in enumerate(tasks):
        ax.barh(yp[i], dur, left=st-.5, height=.6, color=c, alpha=.88,
            edgecolor='white', lw=.5)
        ax.text(st+dur+.2, yp[i], nm, va='center', ha='left', fontsize=10.5,
            color=DRK, fontfamily='Verdana')

    ms = [(4.5,'Prototype Ready'), (9.5,'Experiments Done'),
          (13.5,'Draft Submitted'), (14.8,'Viva Ready')]
    for wk, lb in ms:
        ax.plot(wk, .3, marker='D', ms=12, color=AMB, mec=DRK, mew=.8, zorder=5)
        ax.text(wk, -.5, lb, ha='center', fontsize=9.5, color=DRK,
            fontfamily='Verdana', fontweight='semibold')

    ax.set_xlim(-.2, 16.5)
    ax.set_ylim(-1.3, len(tasks)+1.5)
    ax.set_xlabel('Week', fontsize=12, fontfamily='Verdana', color=DRK, labelpad=12)
    ax.set_xticks(range(1,15))
    ax.set_xticklabels([f'W{i}' for i in range(1,15)], fontsize=10, fontfamily='Verdana')
    ax.set_yticks([])
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.grid(axis='x', alpha=.25, ls='--', color=GRY)

    ax.set_title('Project Schedule — 14-Week Timeline with Milestones',
        fontsize=15, color=DRK, fontweight='bold', fontfamily='Verdana', pad=20)

    leg = [
        mpatches.Patch(color=BLU, alpha=.88, label='Design and Setup'),
        mpatches.Patch(color=STL, alpha=.88, label='Interview Modules'),
        mpatches.Patch(color=TEA, alpha=.88, label='ML and Evaluation'),
        mpatches.Patch(color=GRN, alpha=.88, label='Integration'),
        mpatches.Patch(color=AMB, alpha=.88, label='Writing and Viva'),
        Line2D([0],[0], marker='D', color='w', mfc=AMB, mec=DRK, ms=10, label='Milestone'),
    ]
    ax.legend(handles=leg, loc='upper right', fontsize=10.5, framealpha=.95,
        edgecolor='#e2e8f0', fancybox=True)

    p = os.path.join(OUT, 'final_gantt.png')
    plt.savefig(p, dpi=180, bbox_inches='tight', facecolor=WHT, pad_inches=0.3)
    plt.close()
    print(f'[5/5] Gantt: {p}')
    return p


# ══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_doc(arch, uc, pipe, flow, gantt):
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(2.54)
        s.right_margin = Cm(2.54)

    sty = doc.styles['Normal']
    sty.font.name = 'Arial'
    sty.font.size = Pt(11)
    sty.paragraph_format.line_spacing = 1.5
    sty.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def heading(txt, lv=1):
        h = doc.add_heading(txt, level=lv)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)
            r.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)

    def para(txt, bold=False, italic=False, sa=8):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(txt)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(11); r.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.line_spacing = 1.5

    def figure(path, cap, w=6.0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(cap)
        r.italic = True; r.font.size = Pt(10); r.font.name = 'Arial'
        c.paragraph_format.space_after = Pt(14)

    def bullet(txt):
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(txt)
        r.font.size = Pt(11); r.font.name = 'Arial'

    # ═══════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run('CMP7200 — Individual Masters Project')
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Arial'

    doc.add_paragraph()

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = tp2.add_run('PROJECT PROPOSAL')
    r2.bold = True; r2.font.size = Pt(22); r2.font.name = 'Arial'

    for _ in range(2):
        doc.add_paragraph()

    tp3 = doc.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = tp3.add_run(
        'An Intelligent Multi-Agent AI Interview Platform:\n'
        'Integrating Voice Interaction, Skill Graph Reasoning,\n'
        'and Comparative Answer Evaluation Using\n'
        'LLM Judgement and Trained Machine Learning Models')
    r3.font.size = Pt(13); r3.font.name = 'Arial'

    for _ in range(5):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.line_spacing = 1.8
    for txt, b in [
        ('Birmingham City University\n', True),
        ('Faculty of Computing, Engineering and the Built Environment\n', False),
        ('MSc Computer Science\n', False),
        ('Session 2025-26\n\n', False),
        ('Student Number: ', False),
        ('[INSERT YOUR STUDENT NUMBER HERE]', True),
    ]:
        r = info.add_run(txt)
        r.bold = b; r.font.size = Pt(11); r.font.name = 'Arial'

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════
    heading('Table of Contents', 1)
    toc = [
        'Abstract',
        '1. Background and Rationale',
        '   1.1 Project Aim',
        '   1.2 Objectives',
        '2. Literature Review',
        '   2.1 AI-Driven Recruitment Systems',
        '   2.2 LLMs as Evaluators',
        '   2.3 Knowledge Graphs for Skills Mapping',
        '   2.4 Fairness and Bias in Algorithmic Hiring',
        '   2.5 Conceptual Framework',
        '3. System Overview',
        '   3.1 Technology Stack',
        '   3.2 Use Case Model',
        '   3.3 Interview Process Flow',
        '4. Research Methodology',
        '   4.1 Module 6: Answer Evaluation Pipeline',
        '   4.2 Behavioural Integrity Detection (Module 9)',
        '   4.3 Data Collection and Training Strategy',
        '   4.4 Evaluation Metrics',
        '   4.5 Alternative Approaches Considered',
        '   4.6 Limitations and Mitigation Strategies',
        '   4.7 Ethical Considerations',
        '   4.8 Development Tools and Environment',
        '5. Project Schedule',
        '6. References',
        'Appendix A: Module Summary Table',
    ]
    for t in toc:
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.size = Pt(11); r.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # ABSTRACT
    # ═══════════════════════════════════════════════════
    heading('Abstract', 1)
    para(
        'This proposal outlines the design and planned development of an intelligent '
        'multi-agent AI interview platform. The system combines large language models, '
        'a knowledge-based skill graph, real-time voice interaction, and trained machine '
        'learning classifiers to conduct, monitor, and evaluate candidate interviews. '
        'The platform is structured around twelve modules organised into four phases: '
        'pre-interview setup, live interview session, answer evaluation, and reporting. '
        'The central research contribution is Module 6, which implements two parallel '
        'answer scoring approaches — an LLM-as-Judge method and a trained supervised '
        'classifier using Sentence-BERT embeddings with XGBoost — and compares their '
        'accuracy, consistency, and explainability. An additional trained model (Isolation '
        'Forest) detects behavioural anomalies during interviews. The project follows a '
        'Design Science Research methodology and targets completion within a fourteen-week '
        'schedule. This document presents the project rationale, a critical literature '
        'review, the proposed methodology with alternative approaches considered, and a '
        'detailed project timeline.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 1. BACKGROUND AND RATIONALE
    # ═══════════════════════════════════════════════════
    heading('1. Background and Rationale', 1)

    para(
        'The recruitment industry is going through a major transformation. Over the past '
        'five years, organisations have increasingly adopted AI-powered tools to screen, '
        'interview, and rank job candidates. Platforms like HireVue, Pymetrics, and '
        'myInterview now process millions of candidate assessments each year, frequently '
        'before a human recruiter reviews a single application. The core appeal is '
        'efficiency: faster screening cycles, reduced costs, and more consistent evaluation '
        'across large applicant pools.'
    )

    para(
        'However, this shift has raised serious concerns. In 2019, the Electronic Privacy '
        'Information Centre filed a formal complaint against HireVue, arguing that its use '
        'of facial analysis algorithms in hiring decisions was opaque and potentially '
        'discriminatory (EPIC, 2019). HireVue subsequently discontinued facial analysis in '
        'early 2021, though it continued scoring verbal responses using natural language '
        'processing (HireVue, 2021). The core issue persisted: candidates received scores '
        'but no meaningful explanation of how those scores were calculated.'
    )

    para(
        'The regulatory environment has since caught up. The European Union\'s AI Act, '
        'enacted in 2024, explicitly classifies AI systems used in employment and '
        'recruitment decisions as "high-risk," mandating transparency, human oversight, '
        'and documented bias testing for any such system (European Commission, 2024). '
        'This creates both a challenge and an opportunity for researchers working on '
        'interview automation.'
    )

    para(
        'At the same time, large language models such as GPT-4, Claude, and open-source '
        'alternatives like Llama 3 have reached a capability level where they can hold '
        'natural multi-turn conversations, follow complex scoring rubrics, and provide '
        'structured evaluations. This raises a practical question that forms the basis of '
        'this project: can we build an interview platform that combines the conversational '
        'fluency of LLMs with the measurability and explainability of trained machine '
        'learning classifiers?'
    )

    para(
        'This project proposes exactly that. Rather than building yet another opaque '
        'scoring tool, the aim is to create a system where every evaluation decision is '
        'traceable, where multiple evaluation strategies operate in parallel and can be '
        'compared, and where the gap between what the system scores and what the system '
        'can explain is as narrow as possible.'
    )

    heading('1.1 Project Aim', 2)
    para(
        'To design, build, and evaluate a multi-agent AI interview platform that conducts '
        'voice-based technical interviews, maps candidate skills against job requirements '
        'using a knowledge graph, and evaluates candidate responses through both LLM-based '
        'judgement and a trained machine learning classifier — comparing their accuracy, '
        'consistency, and explainability.'
    )

    heading('1.2 Objectives', 2)
    objs = [
        'Build a skill knowledge graph from CV and job description inputs using NetworkX '
        'and ESCO taxonomy data, mapping candidate competencies against role requirements '
        'and identifying specific skill gaps.',

        'Develop a dynamic question generation agent that adapts interview questions based '
        'on the skill graph structure and the candidate\'s real-time responses, ensuring '
        'coverage of critical skills without redundant questioning.',

        'Implement a voice-based interview agent using OpenAI Whisper for speech-to-text '
        'and a text-to-speech engine for natural conversational interaction, allowing '
        'hands-free interview sessions.',

        'Design and train an answer evaluation classifier using Sentence-BERT embeddings '
        'combined with XGBoost, and compare its scoring performance against an LLM-as-Judge '
        'baseline using inter-rater agreement metrics (Cohen\'s Kappa, Spearman correlation).',

        'Construct a behavioural integrity detection module using an Isolation Forest model '
        'trained on normal interaction patterns — including response timing, tab-switch '
        'frequency, and typing consistency — to flag potentially compromised sessions.',

        'Evaluate the complete system through controlled experiments measuring scoring '
        'agreement with human ratings, consistency under answer paraphrasing, and the '
        'quality of SHAP-based explanations.',
    ]
    for i, o in enumerate(objs, 1):
        p = doc.add_paragraph(style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f'Objective {i}: ')
        r.bold = True; r.font.size = Pt(11); r.font.name = 'Arial'
        r2 = p.add_run(o)
        r2.font.size = Pt(11); r2.font.name = 'Arial'

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 2. LITERATURE REVIEW
    # ═══════════════════════════════════════════════════
    heading('2. Literature Review', 1)

    para(
        'This section critically examines four streams of research that directly shape '
        'the proposed system, before synthesising them into a conceptual framework that '
        'underpins the project design.'
    )

    heading('2.1 AI-Driven Recruitment Systems', 2)
    para(
        'Automated interview systems have evolved rapidly from simple keyword-based resume '
        'screeners to complex multimodal assessment platforms. Hickman et al. (2022) '
        'reviewed text analysis methods used in organisational research and found that '
        'while NLP tools consistently improve scoring reliability compared to unstructured '
        'human evaluation, they frequently lack the criterion validity evidence that '
        'traditional structured interviews have accumulated over decades. In other words, '
        'these tools are consistent but not necessarily accurate.'
    )

    para(
        'Langer et al. (2019) investigated candidate perceptions of AI-driven interviews '
        'and found that applicants perceive automated systems as significantly less fair '
        'than human interviewers, especially when the system provides no feedback or '
        'explanation. This has practical implications: candidates who feel the process was '
        'unfair are less likely to accept offers, undermining the efficiency gains that '
        'motivated the automation in the first place. This perception problem directly '
        'motivates the explainability focus in this project.'
    )

    heading('2.2 LLMs as Evaluators', 2)
    para(
        'Using large language models as automated judges gained significant attention '
        'following the MT-Bench study by Zheng et al. (2023), which demonstrated that '
        'GPT-4 judgements agree with human evaluator preferences at over 80% on open-ended '
        'conversation tasks. However, this promising result must be interpreted carefully. '
        'Stureborg et al. (2024) subsequently identified systematic positional bias in LLM '
        'judges — models tend to favour whichever response appears first, regardless of '
        'actual quality. Wang et al. (2024) found a complementary problem: verbosity bias, '
        'where longer answers receive higher scores even when the additional length adds '
        'no substance.'
    )

    para(
        'The strength of LLM evaluation is its contextual understanding — it can assess '
        'whether an answer genuinely addresses the question. Its weakness is the lack of '
        'transparency and the susceptibility to these systematic biases. A trained '
        'classifier, by contrast, relies on measurable features and can explain its scores '
        'through feature attribution, but may miss nuances that require deeper language '
        'understanding. This trade-off is precisely what Module 6 is designed to investigate.'
    )

    heading('2.3 Knowledge Graphs for Skills Mapping', 2)
    para(
        'Knowledge graphs offer a structured method for representing relationships between '
        'concepts. Chen et al. (2021) applied graph-based models to identify prerequisite '
        'relationships in online learning, showing that graph traversal algorithms can '
        'effectively recommend learning paths. The ESCO framework (European Commission, '
        '2023) provides a standardised taxonomy of over 13,000 skills linked to occupations '
        'across Europe, making it particularly well-suited for recruitment applications '
        'where skills need to be matched against job requirements.'
    )

    para(
        'In this project, the skill graph is not merely decorative. It serves two functional '
        'purposes: gap analysis (comparing what the candidate knows against what the role '
        'requires) and adaptive question targeting (asking questions that probe the edges '
        'of the candidate\'s skill set rather than testing what is already obvious from '
        'their CV). This makes the interview both more efficient and more informative.'
    )

    heading('2.4 Fairness and Bias in Algorithmic Hiring', 2)
    para(
        'Raghavan et al. (2020) audited several commercial hiring algorithms and found that '
        'most lacked any formal fairness testing protocol. Bogen and Rieke (2018) warned '
        'that algorithmic tools risk encoding historical biases from training data, '
        'particularly around gender, ethnicity, and socioeconomic background. The EU AI Act '
        'now requires documented bias testing for high-risk systems (European Commission, '
        '2024), creating regulatory pressure for the kind of transparency this project aims '
        'to provide.'
    )

    para(
        'This project addresses fairness at two levels. The trained classifier uses '
        'exclusively skill-based and linguistic features — semantic similarity, keyword '
        'coverage, grammar quality, response structure — rather than features that could '
        'serve as demographic proxies. Additionally, the SHAP explanation layer enables '
        'auditing of whether any feature disproportionately influences scores for '
        'particular candidate groups. This does not guarantee fairness, but it makes '
        'unfairness detectable.'
    )

    heading('2.5 Conceptual Framework', 2)
    para(
        'Synthesising the four research themes above, a clear conceptual model emerges. '
        'Current AI interview tools suffer from three interconnected weaknesses: opacity '
        '(they cannot explain their decisions), bias susceptibility (they may encode '
        'historical discrimination), and single-method fragility (they rely on one '
        'evaluation approach with no cross-validation). This project addresses all three '
        'through a multi-agent architecture where: (a) explainability is built into the '
        'evaluation pipeline via SHAP, (b) fairness is promoted through skill-based features '
        'and auditability, and (c) robustness is achieved through dual-track scoring where '
        'two independent evaluation methods serve as mutual checks. The skill graph adds a '
        'fourth dimension — structured knowledge representation — that grounds the interview '
        'process in a documented competency framework rather than ad-hoc question selection.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 3. SYSTEM OVERVIEW
    # ═══════════════════════════════════════════════════
    heading('3. System Overview', 1)

    para(
        'The platform consists of twelve modules arranged into four sequential phases. '
        'Each module is designed as an independent agent with defined inputs and outputs, '
        'allowing them to be developed, tested, and if necessary replaced independently. '
        'Figure 1 presents the complete system architecture.'
    )

    figure(arch, 'Figure 1: System Architecture — Four-Phase Layered Design', w=6.5)

    heading('3.1 Technology Stack', 2)
    para(
        'Table 1 summarises the core technology behind each module and indicates whether '
        'the component is pretrained, trained by the researcher, or rule-based. Two modules '
        '(M6 and M9) involve model training — these form the primary research contribution.'
    )

    tbl = doc.add_table(rows=13, cols=4)
    tbl.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Module', 'Function', 'Core Technology', 'Type']):
        c = tbl.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10); r.font.name = 'Arial'

    data = [
        ('M1', 'CV Parsing', 'LLM (GPT-4 / Llama 3)', 'Pretrained'),
        ('M2', 'JD Understanding', 'LLM (same)', 'Pretrained'),
        ('M3', 'Skill Graph', 'NetworkX + ESCO Taxonomy', 'Rule-based'),
        ('M4', 'Question Generation', 'LLM + Graph Traversal', 'Pretrained'),
        ('M5', 'Voice Interview', 'Whisper STT + Google TTS', 'Pretrained'),
        ('M6', 'Answer Evaluation', 'LLM + S-BERT + XGBoost', 'Trained'),
        ('M7', 'Vision Monitor', 'MediaPipe Face Mesh', 'Pretrained'),
        ('M8', 'Posture Analysis', 'MediaPipe Pose', 'Pretrained'),
        ('M9', 'Behavioural Integrity', 'Isolation Forest', 'Trained'),
        ('M10', 'Emotion Detection', 'wav2vec2-emotion', 'Pretrained'),
        ('M11', 'Recommendation', 'Weighted Fusion + LLM', 'Hybrid'),
        ('M12', 'Report Generation', 'Template Engine', 'Rule-based'),
    ]
    for ri, (m, f, t, tp) in enumerate(data, 1):
        for ci, v in enumerate([m, f, t, tp]):
            cell = tbl.rows[ri].cells[ci]
            cell.text = v
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10); r.font.name = 'Arial'

    tc = doc.add_paragraph()
    tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tc.add_run('Table 1: Module Technology Stack — Trained vs Pretrained Components')
    r.italic = True; r.font.size = Pt(10); r.font.name = 'Arial'
    tc.paragraph_format.space_after = Pt(14)

    heading('3.2 Use Case Model', 2)
    para(
        'Two primary actors interact with the system. The Candidate provides their CV, '
        'participates in the voice interview, and answers technical questions. The Recruiter '
        'receives the evaluation report, reviews flagged answers, and makes the final hiring '
        'decision. Figure 2 shows these interactions.'
    )
    figure(uc, 'Figure 2: Use Case Diagram — Candidate and Recruiter Interactions', w=6.0)

    heading('3.3 Interview Process Flow', 2)
    para(
        'Figure 3 shows how data moves through the system during a complete interview '
        'session, from initial CV upload through skill mapping, voice interview, evaluation, '
        'and final reporting.'
    )
    figure(flow, 'Figure 3: Data Flow Diagram — End-to-End Interview Process', w=6.2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 4. RESEARCH METHODOLOGY
    # ═══════════════════════════════════════════════════
    heading('4. Research Methodology', 1)

    para(
        'This project follows a Design Science Research (DSR) methodology as formalised '
        'by Hevner et al. (2004). DSR is specifically intended for research that creates '
        'and evaluates IT artefacts, making it well-suited for a project that builds a '
        'working software system and then evaluates its performance. The approach involves '
        'iterative cycles of design, implementation, and evaluation. The central research '
        'question is:'
    )

    rq = doc.add_paragraph()
    rq.paragraph_format.left_indent = Cm(1.5)
    rq.paragraph_format.right_indent = Cm(1.5)
    r = rq.add_run(
        'RQ: How does a trained supervised classifier (Sentence-BERT + XGBoost) compare '
        'to an LLM-as-Judge approach for evaluating interview responses, in terms of '
        'agreement with human ratings, scoring consistency under paraphrasing, and '
        'explainability of decisions?')
    r.italic = True; r.font.size = Pt(11); r.font.name = 'Arial'
    rq.paragraph_format.space_after = Pt(14)

    heading('4.1 Module 6 — Answer Evaluation Pipeline (Core Contribution)', 2)
    para(
        'Module 6 represents the primary research contribution. It implements two '
        'independent evaluation strategies that process every candidate answer in parallel, '
        'then compares their outputs to determine where they agree and where they diverge. '
        'Figure 4 illustrates this dual-track pipeline.'
    )

    figure(pipe, 'Figure 4: Module 6 — Dual-Track Answer Evaluation Pipeline', w=6.5)

    heading('4.1.1 Track A: LLM-as-Judge', 3)
    para(
        'The candidate\'s transcribed answer is sent to a large language model (GPT-4 or '
        'Llama 3) along with a structured rubric prompt specifying four scoring criteria: '
        'technical accuracy, completeness, clarity, and relevance. The model returns a '
        'numerical score between 0 and 100 with a brief justification. To mitigate the '
        'positional bias documented by Stureborg et al. (2024), the system randomises '
        'prompt element ordering across multiple calls and averages the results.'
    )

    heading('4.1.2 Track B: Trained ML Classifier', 3)
    para(
        'The second approach converts each answer into a fixed set of measurable features '
        'before classification. The feature set includes:'
    )
    for f in [
        'Sentence-BERT embedding (384-dimensional semantic vector)',
        'Cosine similarity between answer and ideal reference answer',
        'Keyword coverage score (proportion of expected technical terms)',
        'Grammar and fluency score (via language-tool-python)',
        'Response length and structural metrics (sentences, paragraphs)',
        'Specificity score (ratio of concrete terms to vague fillers)',
    ]:
        bullet(f)

    para(
        'These features are input to an XGBoost classifier trained on labelled interview '
        'data. For every prediction, SHAP values are computed to identify which features '
        'contributed most to the score, providing full transparency into the decision.'
    )

    heading('4.1.3 Comparison and Fusion Strategy', 3)
    para(
        'Both tracks independently process every answer. The system compares outputs using '
        'Cohen\'s Kappa (categorical agreement), Pearson correlation (numerical alignment), '
        'and disagreement threshold analysis (flagging cases where scores differ by more '
        'than 20 points). Where both methods agree, the system reports high confidence. '
        'Where they substantially disagree, the answer is flagged for human review.'
    )

    heading('4.2 Behavioural Integrity Detection (Module 9)', 2)
    para(
        'Module 9 uses an Isolation Forest — an unsupervised anomaly detection algorithm — '
        'trained on normal interview interaction patterns. Input features include response '
        'timing distributions, tab-switch frequency, browser focus events, mouse movement '
        'entropy, typing speed consistency, and speech hesitation patterns. Sessions that '
        'deviate significantly from the learned baseline are flagged as potentially '
        'compromised. This module does not make hiring decisions; it provides a confidence '
        'annotation that the recruiter can interpret in context.'
    )

    heading('4.3 Data Collection and Training Strategy', 2)
    para(
        'The trained classifier requires labelled interview answer data. Three data sources '
        'will be combined. First, publicly available interview Q&A datasets from Kaggle and '
        'HuggingFace that include quality ratings. Second, synthetic training data generated '
        'by prompting LLMs to produce answers at controlled quality levels (strong, medium, '
        'weak) — an established augmentation strategy validated by Ye et al. (2022). Third, '
        'a manually-labelled validation set of approximately 200 answers rated by two '
        'independent raters, used exclusively for testing to measure real-world performance.'
    )

    para(
        'For the Isolation Forest, baseline data will come from pilot sessions where '
        'volunteer participants complete mock interviews under normal conditions, '
        'establishing the distribution that anomalies are measured against.'
    )

    heading('4.4 Evaluation Metrics', 2)
    para('The system will be evaluated across three dimensions:')

    para(
        'Scoring Accuracy — How well does each method agree with human judgement? Measured '
        'using quadratic-weighted Cohen\'s Kappa and Spearman rank correlation against the '
        'manually-labelled validation set.', sa=4)
    para(
        'Consistency — Does the system produce stable scores for semantically equivalent '
        'answers? Tested by paraphrasing 50 answers and measuring score variance across '
        'each paraphrase set.', sa=4)
    para(
        'Explainability — Can the system meaningfully justify its scores? Evaluated '
        'qualitatively by checking whether SHAP feature attributions align with expert '
        'intuition about what makes an answer strong or weak.', sa=12)

    heading('4.5 Alternative Approaches Considered', 2)
    para(
        'Several alternative approaches were considered before settling on the proposed '
        'design. For answer evaluation, one option was to use only the LLM-as-Judge '
        'approach, which would be simpler to implement. However, this would provide no '
        'independent validation and would inherit all the biases documented in the '
        'literature (Stureborg et al., 2024; Wang et al., 2024). Another option was to '
        'train a deep neural network (such as a fine-tuned BERT classifier) instead of '
        'XGBoost. While this could potentially achieve higher accuracy, it would sacrifice '
        'the interpretability that SHAP provides for tree-based models, which is central '
        'to the project\'s explainability goals.'
    )

    para(
        'For the skill graph, a graph neural network (GNN) approach was considered but '
        'rejected on the grounds that the ESCO taxonomy is already well-structured and '
        'does not require learned embeddings — a deterministic graph traversal is both '
        'simpler and more transparent. For voice interaction, a fully custom ASR model '
        'was briefly considered but dismissed in favour of Whisper, which already achieves '
        'near-human accuracy on English speech and would divert effort from the core '
        'research question.'
    )

    heading('4.6 Limitations and Mitigation Strategies', 2)
    para(
        'Several limitations are acknowledged. First, the trained classifier\'s performance '
        'depends on the quality and diversity of training data. If the available interview '
        'datasets are too narrow in topic coverage, the classifier may not generalise well '
        'to unseen question domains. This will be mitigated by using synthetic data '
        'augmentation across a wide range of technical topics and by clearly documenting '
        'the model\'s performance boundaries.'
    )

    para(
        'Second, the LLM-as-Judge approach relies on API access to commercial models, '
        'which introduces cost and latency constraints. This will be managed by caching '
        'all API responses and budgeting approximately £100 for experiment runs, with '
        'open-source Llama 3 as a fallback.'
    )

    para(
        'Third, the multimodal sensor modules (M7, M8, M10) are secondary to the core '
        'evaluation research. If they prove too time-consuming to implement fully, they '
        'will be simplified to basic implementations while Module 6 receives priority. '
        'The modular architecture explicitly supports this: each module can be developed '
        'or deferred independently without affecting the others.'
    )

    para(
        'Fourth, the behavioural integrity model (M9) is trained on pilot data that may '
        'not represent the full range of real-world interview behaviours. This limitation '
        'will be explicitly discussed in the evaluation chapter, and the model\'s false '
        'positive rate will be reported alongside its detection accuracy.'
    )

    heading('4.7 Ethical Considerations', 2)
    para(
        'This project does not involve real hiring decisions and will not affect any '
        'person\'s employment prospects. All evaluation experiments use synthetic or '
        'publicly available data. If pilot testing involves human participants for Module 9 '
        'baseline collection, BCU ethical approval will be obtained before any data '
        'collection begins. Participants will give informed written consent, their data '
        'will be anonymised and stored securely in accordance with GDPR, and they will '
        'retain the right to withdraw at any point without consequence. The system is '
        'designed as a research prototype, not a production hiring tool.'
    )

    heading('4.8 Development Tools and Environment', 2)
    para(
        'The system will be built using Python 3.11 with FastAPI as the backend framework. '
        'The frontend will use Streamlit for rapid prototyping, with the option to migrate '
        'to React if time permits. Key libraries include NetworkX for graph operations, '
        'HuggingFace Transformers for Sentence-BERT and Whisper, XGBoost and SHAP for the '
        'trained classifier, scikit-learn for evaluation metrics, and MediaPipe for the '
        'vision and posture modules. Source code will be version-controlled using Git. '
        'Experiment tracking will use Weights and Biases to ensure full reproducibility '
        'of model training runs and hyperparameter searches.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 5. PROJECT SCHEDULE
    # ═══════════════════════════════════════════════════
    heading('5. Project Schedule', 1)

    para(
        'The project spans fourteen weeks across five overlapping phases: literature '
        'review and setup, core module development, ML model training and experimentation, '
        'system integration, and dissertation writing. Figure 5 presents the schedule with '
        'four key milestones.'
    )

    figure(gantt, 'Figure 5: 14-Week Project Schedule with Milestones', w=6.5)

    para(
        'Milestone 1 (Week 4-5): A working prototype demonstrating the interview flow '
        'from CV upload through question generation to voice interaction. Milestone 2 '
        '(Week 9-10): Both evaluation models trained and comparison experiments completed. '
        'Milestone 3 (Week 13): Final dissertation draft submitted to supervisor. '
        'Milestone 4 (Week 14): Viva presentation prepared and rehearsed.'
    )

    para(
        'The schedule includes deliberate buffer time in weeks 10-11 for unexpected delays '
        'in model training or data preparation. Dissertation writing overlaps with the '
        'later implementation phases to avoid a compressed writing period at the end. If '
        'the multimodal sensor modules require more time than anticipated, they will be '
        'descoped to basic implementations while the core Module 6 research is prioritised.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 6. REFERENCES
    # ═══════════════════════════════════════════════════
    heading('References', 1)

    refs = [
        'Bogen, M. and Rieke, A. (2018) Help Wanted: An Examination of Hiring '
        'Algorithms, Equity, and Bias. Washington, DC: Upturn.',

        'Chen, Y., Li, X. and Zhang, J. (2021) \'A knowledge graph approach to '
        'prerequisite identification in online learning\', Computers and Education: '
        'Artificial Intelligence, 2(1), pp. 1-12. doi: 10.1016/j.caeai.2021.100016.',

        'EPIC (2019) Complaint and Request for Investigation of HireVue, Inc. '
        'Washington, DC: Electronic Privacy Information Centre.',

        'European Commission (2023) ESCO: European Skills, Competences, Qualifications '
        'and Occupations. Available at: https://esco.ec.europa.eu/ '
        '(Accessed: 15 March 2026).',

        'European Commission (2024) Regulation (EU) 2024/1689 laying down harmonised '
        'rules on artificial intelligence (Artificial Intelligence Act). Official '
        'Journal of the European Union, L series.',

        'Hevner, A.R., March, S.T., Park, J. and Ram, S. (2004) \'Design science in '
        'information systems research\', MIS Quarterly, 28(1), pp. 75-105.',

        'Hickman, L., Thapa, S., Tay, L., Cao, M. and Srinivasan, P. (2022) \'Text '
        'preprocessing for text mining in organizational research: Review and '
        'recommendations\', Organizational Research Methods, 25(1), pp. 114-146.',

        'HireVue (2021) HireVue Leads the Industry with Removal of Visual Analysis '
        'from Assessments. Press Release, 11 January.',

        'Langer, M., Konig, C.J. and Papathanasiou, M. (2019) \'Highly automated '
        'job interviews: Acceptance under the influence of stakes\', International '
        'Journal of Selection and Assessment, 27(3), pp. 217-234.',

        'Raghavan, M., Barocas, S., Kleinberg, J. and Levy, K. (2020) \'Mitigating '
        'bias in algorithmic hiring: Evaluating claims and practices\', in Proceedings '
        'of the 2020 ACM Conference on Fairness, Accountability, and Transparency. '
        'New York: ACM, pp. 469-481.',

        'Stureborg, R., Alikaniotis, D. and Suhara, Y. (2024) \'Large language models '
        'are inconsistent and biased evaluators\', arXiv preprint, arXiv:2405.01724.',

        'Wang, P., Li, L., Chen, L. and Song, D. (2024) \'Large language models are '
        'not fair evaluators\', in Proceedings of the 62nd Annual Meeting of the '
        'Association for Computational Linguistics. ACL.',

        'Ye, J., Chen, J., Liu, Q., Xu, Z. and Wan, X. (2022) \'Generative data '
        'augmentation for commonsense reasoning\', in Findings of the Association for '
        'Computational Linguistics: EMNLP 2022, pp. 1008-1025.',

        'Zheng, L., Chiang, W.L., Sheng, Y., Zhuang, S., Wu, Z., Zhuang, Y., Lin, Z., '
        'Li, Z., Li, D., Xing, E., Zhang, H., Gonzalez, J.E. and Stoica, I. (2023) '
        '\'Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena\', Advances in Neural '
        'Information Processing Systems, 36.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        r.font.size = Pt(11); r.font.name = 'Arial'
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # APPENDIX A
    # ═══════════════════════════════════════════════════
    heading('Appendix A: Module Summary', 1)

    para(
        'Table A1 provides a complete summary of all twelve modules, their phase '
        'allocation, and their role within the system.'
    )

    atbl = doc.add_table(rows=13, cols=5)
    atbl.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Module', 'Phase', 'Role', 'Technology', 'Training']):
        c = atbl.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9); r.font.name = 'Arial'

    app_data = [
        ('M1', '1', 'Extract skills and experience from CV', 'LLM', 'No'),
        ('M2', '1', 'Parse job requirements and skill needs', 'LLM', 'No'),
        ('M3', '1', 'Map skills into knowledge graph', 'NetworkX', 'No'),
        ('M4', '1', 'Generate adaptive interview questions', 'LLM + Graph', 'No'),
        ('M5', '2', 'Conduct voice-based interview', 'Whisper + TTS', 'No'),
        ('M6', '3', 'Score answers (dual-track evaluation)', 'LLM + XGBoost', 'Yes'),
        ('M7', '2', 'Track facial attention and gaze', 'MediaPipe', 'No'),
        ('M8', '2', 'Analyse body language and posture', 'MediaPipe', 'No'),
        ('M9', '2', 'Detect anomalous interview behaviour', 'Isolation Forest', 'Yes'),
        ('M10', '2', 'Detect emotional state from voice', 'wav2vec2', 'No'),
        ('M11', '4', 'Fuse scores into recommendation', 'Rules + LLM', 'No'),
        ('M12', '4', 'Generate final candidate report', 'Templates', 'No'),
    ]
    for ri, row_data in enumerate(app_data, 1):
        for ci, v in enumerate(row_data):
            cell = atbl.rows[ri].cells[ci]
            cell.text = v
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = 'Arial'

    tc2 = doc.add_paragraph()
    tc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tc2.add_run('Table A1: Complete Module Summary')
    r.italic = True; r.font.size = Pt(10); r.font.name = 'Arial'

    # Save
    path = os.path.join(OUT, 'AI_Interview_Final_Proposal.docx')
    doc.save(path)
    print(f'[DOC] {path}')
    return path


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    a = build_architecture()
    u = build_usecase()
    p = build_pipeline()
    f = build_dataflow()
    g = build_gantt()
    build_doc(a, u, p, f, g)
    print('\n=== All files generated successfully ===')
